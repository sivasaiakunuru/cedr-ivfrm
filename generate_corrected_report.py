#!/usr/bin/env python3
"""
Generate CYB400 Automobility Cybersecurity Strategy Final Report - Corrected Version
Topic: Securing Over-the-Air (OTA) Updates
Corrections: Improved formatting, consistent citations, better tables
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell):
    """Set cell border properties"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_elm = OxmlElement(f'w:{edge}')
        edge_elm.set(qn('w:val'), 'single')
        edge_elm.set(qn('w:sz'), '4')
        edge_elm.set(qn('w:color'), '000000')
        tcBorders.append(edge_elm)
    tcPr.append(tcBorders)

def format_table(table):
    """Format table with consistent styling"""
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

def add_heading_custom(doc, text, level):
    """Add a heading with consistent formatting"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        else:
            run.font.size = Pt(12)
            run.font.bold = True
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, first_line_indent=0.3):
    """Add a paragraph with custom formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet_list(doc, items):
    """Add a bulleted list"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)

def create_title_page(doc):
    """Create the title page"""
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Automobility Cybersecurity Strategy:\nSecuring Over-the-Air (OTA) Updates")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Final Project Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("Course: CYB400 – Automobility Cybersecurity Strategy")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("Zekelman School of Business | St. Clair College")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    group = doc.add_paragraph()
    group.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = group.add_run("Group 8")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    
    members_text = "Siva Sai Akunuru | Manish Reddy Ennedla\nSri Venkata Naga Sai Chinmai Malladi | Lotachi Obi-Nwaigwe"
    members = doc.add_paragraph()
    members.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = members.add_run(members_text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    prof = doc.add_paragraph()
    prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prof.add_run("Professor: A. Sodiq Shofoluwe")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("Date: April 2026")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def create_toc(doc):
    """Create table of contents"""
    add_heading_custom(doc, "Table of Contents", 1)
    
    toc_items = [
        ("Executive Summary", 2),
        ("1. Introduction and Background", 4),
        ("2. System Description and Operating Context", 6),
        ("3. Security Objectives and Scope", 8),
        ("4. Ethical and Societal Impact Assessment", 10),
        ("5. Threat Modeling and Risk Analysis", 12),
        ("6. Security Architecture and Technical Controls", 15),
        ("7. Process and Governance Model", 18),
        ("8. Business and Strategic Considerations", 20),
        ("9. Validation, Feasibility, and Limitations", 22),
        ("10. Conclusion and Future Considerations", 24),
        ("References", 26),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        
        # Add page number aligned right
        p.add_run("\t\t\t\t\t\t\t\t\t" + str(page))
        p.runs[1].font.name = 'Times New Roman'
        p.runs[1].font.size = Pt(11)
    
    doc.add_page_break()

def create_executive_summary(doc):
    """Create executive summary section"""
    add_heading_custom(doc, "Executive Summary", 1)
    
    paras = [
        "The automotive industry is undergoing a foundational transformation as vehicles evolve from mechanical systems to software-defined platforms. Over-the-Air (OTA) updates enable manufacturers to deploy firmware patches, security fixes, and feature enhancements remotely without requiring dealership visits. Industry projections indicate that by 2030, over 80% of vehicle software will be distributed via OTA channels, making OTA security foundational to modern automobility (McKinsey & Company, 2023).",
        
        "This report presents a comprehensive cybersecurity strategy for securing OTA update systems in connected vehicles. The strategy addresses the complete ecosystem encompassing OEM cloud infrastructure, telematics control units (TCUs), electronic control units (ECUs), in-vehicle networks, and wireless communication channels. The security architecture employs defense-in-depth principles, integrating cryptographic authenticity verification, encrypted transmission protocols, secure boot validation, and anomaly detection mechanisms.",
        
        "Key security objectives include: (1) ensuring authenticity and integrity through ECDSA cryptographic signing with Hardware Security Module (HSM) backing; (2) protecting confidentiality via TLS 1.3 encryption; (3) preventing replay and rollback attacks through nonce-based freshness verification and monotonic version counters; (4) detecting anomalous behavior using lightweight ML-based intrusion detection; and (5) enabling safe recovery through automated rollback to known-good firmware.",
        
        "The strategy demonstrates alignment with ISO/SAE 21434 for automotive cybersecurity engineering, UNECE WP.29 R155 for vehicle cybersecurity regulation, and NIST SP 800-193 for platform firmware resiliency. The ethical framework prioritizes human safety and privacy through privacy-by-design principles, user consent mechanisms, and fail-safe protections. The business case demonstrates that security investment is justified by regulatory compliance, avoided recall costs exceeding $150 million for major incidents such as the 2015 Jeep Cherokee recall, competitive differentiation, and enablement of emerging software-defined vehicle revenue models.",
        
        "This strategy provides a defensible, implementable framework for securing OTA updates that balances technical rigor, ethical responsibility, and business practicality in the rapidly evolving landscape of connected and autonomous vehicles."
    ]
    
    for para in paras:
        add_paragraph_custom(doc, para)

def create_introduction(doc):
    """Create introduction section"""
    add_heading_custom(doc, "1. Introduction and Background", 1)
    add_heading_custom(doc, "1.1 The Software-Defined Vehicle Revolution", 2)
    
    add_paragraph_custom(doc, "The automotive industry is undergoing foundational transformation. Modern vehicles are increasingly defined by software governing Advanced Driver-Assistance Systems (ADAS), powertrain management, infotainment, and autonomous driving functions. Contemporary premium vehicles contain over 100 million lines of code across 100+ Electronic Control Units (ECUs), with complexity accelerating annually (McKinsey & Company, 2023). This evolution enables the software-defined vehicle paradigm, where functionality can be continuously enhanced post-manufacturing through OTA updates.")
    
    add_heading_custom(doc, "1.2 The OTA Update Paradigm", 2)
    add_paragraph_custom(doc, "OTA updates encompass Firmware-Over-the-Air (FOTA) for ECU firmware and Software-Over-the-Air (SOTA) for application-layer software. Updates leverage cellular (4G/5G) and Wi-Fi to transmit packages from OEM cloud infrastructure to vehicle fleets. The process involves: (1) package creation and cryptographic signing at the OEM backend; (2) cloud distribution via Content Delivery Networks; (3) encrypted wireless transmission; (4) TCU reception and verification; (5) ECU distribution via Controller Area Network (CAN) bus; and (6) post-installation validation (ISO/SAE, 2021).")
    
    add_heading_custom(doc, "1.3 The Cybersecurity Imperative", 2)
    add_paragraph_custom(doc, "OTA capability creates critical attack surfaces exploitable at scale. In 2015, security researchers Charlie Miller and Chris Valasek demonstrated remote exploitation of a Jeep Cherokee, gaining wireless control over steering, braking, and transmission systems through connectivity vulnerabilities. This demonstration led to the recall of 1.4 million vehicles and served as a decisive wake-up call for the automotive industry (Miller & Valasek, 2015). Automotive cybersecurity incidents continue increasing, with remote attacks constituting the majority and OTA vulnerabilities representing a growing concern (Upstream Security, 2023).")
    
    add_heading_custom(doc, "1.4 Regulatory and Standards Landscape", 2)
    add_paragraph_custom(doc, "In response to mounting threats, regulators have established comprehensive frameworks. The United Nations Economic Commission for Europe (UNECE) WP.29 Regulation No. 155 mandates that vehicle manufacturers implement a certified Cybersecurity Management System (CSMS) and demonstrate cybersecurity risk management throughout the vehicle lifecycle (UNECE, 2020). ISO/SAE 21434 provides a comprehensive engineering standard for automotive cybersecurity covering risk assessment, security design, verification, and validation (ISO/SAE, 2021). NIST SP 800-193 offers guidance on platform firmware resiliency, addressing the protection, detection, and recovery of firmware from unauthorized modification (NIST, 2018).")

def create_system_description(doc):
    """Create system description section"""
    add_heading_custom(doc, "2. System Description and Operating Context", 1)
    add_heading_custom(doc, "2.1 OTA Ecosystem Architecture", 2)
    
    add_paragraph_custom(doc, "The OTA update ecosystem is a distributed, multi-layered system spanning cloud infrastructure, wireless communication networks, and in-vehicle computing platforms. A thorough understanding of this architecture is essential for identifying vulnerabilities and designing effective security controls (ISO/SAE, 2021).")
    
    add_heading_custom(doc, "2.2 Cloud-Side Components", 2)
    items = [
        "Build and Signing Server: The OEM's secure build environment where firmware packages are compiled, tested, and cryptographically signed using private keys stored in Hardware Security Modules (HSMs). The signing process ensures only authorized, verified software enters the distribution pipeline (NIST, 2018).",
        
        "Update Management Platform: The orchestration layer responsible for managing update campaigns, including targeting specific vehicle models, firmware versions, and geographic regions. This platform manages phased rollouts, monitors deployment progress, and triggers rollbacks if anomalies are detected (ISO/SAE, 2021).",
        
        "MQTT/HTTPS Broker: The message broker facilitating communication between the cloud backend and vehicles. MQTT (Message Queuing Telemetry Transport) is widely used in automotive OTA systems due to its lightweight publish-subscribe architecture, which is well-suited for constrained vehicle environments (Halder et al., 2020).",
        
        "Content Delivery Network (CDN): Distributes update packages to geographically distributed vehicle populations, ensuring low-latency, high-availability delivery."
    ]
    add_bullet_list(doc, items)
    
    add_heading_custom(doc, "2.3 Vehicle-Side Components", 2)
    items = [
        "Telematics Control Unit (TCU): The vehicle's primary connectivity module, responsible for receiving update packages from the cloud, performing initial validation, and coordinating distribution to target ECUs. The TCU serves as the gateway between the external network and the internal vehicle network (Upstream Security, 2023).",
        
        "Gateway ECU: Acts as a firewall and traffic controller for the in-vehicle network, enforcing access control policies and filtering communications between different network domains to separate safety-critical powertrain ECUs from infotainment systems (Checkoway et al., 2011).",
        
        "Target ECUs: The individual electronic control units that receive and install firmware updates, including safety-critical systems (ADAS, braking, steering), body control modules, powertrain controllers, and infotainment units.",
        
        "Hardware Security Module (HSM): A tamper-resistant hardware component that stores cryptographic keys, performs signature verification, and supports secure boot processes. The HSM serves as the vehicle's hardware root of trust for OTA update verification (NIST, 2018)."
    ]
    add_bullet_list(doc, items)

def create_security_objectives(doc):
    """Create security objectives section"""
    add_heading_custom(doc, "3. Security Objectives and Scope", 1)
    add_heading_custom(doc, "3.1 Security Objectives", 2)
    
    objectives = [
        "Authenticity and Integrity Assurance: Ensure that every OTA update package received by a vehicle is authentic and unmodified. This is achieved through ECDSA cryptographic code signing with keys anchored in HSMs at both cloud and vehicle sides (NIST, 2018).",
        
        "Confidentiality of Update Payloads: Protect the content of OTA update packages from unauthorized disclosure during transmission using TLS 1.3 encryption and at-rest encryption on the server side.",
        
        "Replay and Rollback Prevention: Prevent adversaries from retransmitting previously captured update packages or forcing reversion to vulnerable firmware versions through nonce-based freshness verification and monotonic version counters stored in the HSM.",
        
        "Anomaly Detection and Monitoring: Detect anomalous or malicious behavior associated with OTA processes at the vehicle and fleet levels using lightweight machine-learning-based intrusion detection systems.",
        
        "Safe Recovery and Resilience: Ensure vehicles can safely recover from failed or compromised updates through automated rollback, A/B partition schemes, and fail-safe modes that maintain essential vehicle functionality (NIST, 2018)."
    ]
    
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(obj)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
    
    add_heading_custom(doc, "3.2 Scope and Exclusions", 2)
    
    # Scope table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'In-Scope Elements'
    hdr_cells[1].text = 'Exclusions'
    
    scope_data = [
        ("OEM cloud infrastructure (build servers, signing infrastructure, update platform)", "Physical hardware root-of-trust at the silicon fabrication level"),
        ("Cellular (4G/5G) and Wi-Fi communication channels", "Non-OTA update mechanisms (USB flashing, dealership diagnostic tools)"),
        ("TCU, Gateway ECU, target ECUs, HSM", "Vehicle-to-Everything (V2X) communications security"),
        ("CAN bus and automotive Ethernet networks", "Autonomous driving algorithm functional safety validation"),
        ("OEMs, Tier-1/Tier-2 suppliers, vehicle owners, regulators", "Legacy vehicles without connectivity or HSM capability")
    ]
    
    for in_scope, exclusion in scope_data:
        row_cells = table.add_row().cells
        row_cells[0].text = in_scope
        row_cells[1].text = exclusion
    
    format_table(table)

def create_ethical_assessment(doc):
    """Create ethical assessment section"""
    add_heading_custom(doc, "4. Ethical and Societal Impact Assessment", 1)
    
    add_paragraph_custom(doc, "The security of OTA update mechanisms carries profound ethical implications. Vehicles are safety-critical systems operating in shared public spaces, carrying human occupants, and interacting with pedestrians, cyclists, and other road users (UNECE, 2020). This section examines the ethical dimensions of OTA cybersecurity and describes how the proposed strategy addresses each concern, grounded in principles from the NIST Privacy Framework (NIST, 2020) and UNECE WP.29 regulatory requirements.")
    
    add_heading_custom(doc, "4.1 Safety and the Duty of Care", 2)
    add_paragraph_custom(doc, "The most fundamental ethical imperative is the protection of human life. Original Equipment Manufacturers (OEMs) bear a duty of care to ensure their vehicles do not pose unreasonable risks. When OTA updates modify safety-critical systems such as ADAS algorithms, electronic stability control, or braking calibration, the consequences of a compromised update are potentially fatal (Miller & Valasek, 2015). The strategy response includes cryptographic signature verification ensuring only authorized firmware is installed on safety-critical ECUs, dual-bank A/B partitioning guaranteeing a known-good firmware image is always available for fallback, fail-safe modes maintaining essential vehicle functions even if an update fails, and phased rollout policies deploying safety-critical updates to a canary group before wider deployment.")
    
    add_heading_custom(doc, "4.2 Privacy and Data Protection", 2)
    add_paragraph_custom(doc, "The OTA process inherently involves collecting vehicle data including Vehicle Identification Numbers (VINs), firmware versions, geographic location, and connectivity status. If improperly handled, this data could enable tracking, profiling, or surveillance of vehicle owners (NIST, 2020). Strategy responses include data minimization collecting only strictly necessary information with telemetry anonymized before cloud transmission, informed consent requiring explicit vehicle owner consent where required by regulation, TLS 1.3 encryption for all data in transit, retention limits minimizing personal data storage periods, and regulatory compliance with GDPR, CCPA, and the NIST Privacy Framework.")
    
    add_heading_custom(doc, "4.3 Fairness, Transparency, and Societal Trust", 2)
    add_paragraph_custom(doc, "A critical ethical concern is the potential for OTA updates to create or exacerbate inequities if security patches are prioritized based on vehicle tier, geography, or subscription status (Lim et al., 2022). The strategy mandates that security-critical patches are distributed to all affected vehicles regardless of model, subscription status, or location. Transparency requirements include clear release notes describing changes in accessible language, publicly accessible cybersecurity disclosure policies, and timely incident notification. These measures maintain societal trust in connected mobility essential for widespread technology adoption (Upstream Security, 2023).")

def create_threat_modeling(doc):
    """Create threat modeling section"""
    add_heading_custom(doc, "5. Threat Modeling and Risk Analysis", 1)
    add_heading_custom(doc, "5.1 Methodology and Adversary Profiles", 2)
    
    add_paragraph_custom(doc, "The threat modeling process employs the STRIDE framework (Spoofing, Tampering, Repudiation, Information Disclosure, Denial of Service, Elevation of Privilege), adapted for the automotive OTA context and supplemented by the ENISA threat taxonomy for smart vehicles (ENISA, 2019). This approach aligns with the Threat Analysis and Risk Assessment (TARA) methodology specified in ISO/SAE 21434, Clause 15 (ISO/SAE, 2021).")
    
    # Adversary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Adversary Type'
    hdr_cells[1].text = 'Motivation'
    hdr_cells[2].text = 'Capability'
    
    adversaries = [
        ("Script Kiddies", "Curiosity, notoriety", "Low; publicly available tools"),
        ("Organized Criminals", "Financial gain", "Moderate–High; custom malware, ransomware"),
        ("Nation-State Actors", "Espionage, sabotage", "Very High; APT capabilities, zero-days (ENISA, 2019)"),
        ("Malicious Insiders", "Financial gain, revenge", "High; privileged access"),
        ("Security Researchers", "Public safety, recognition", "High; deep technical expertise (Miller & Valasek, 2015)")
    ]
    
    for adv, mot, cap in adversaries:
        row_cells = table.add_row().cells
        row_cells[0].text = adv
        row_cells[1].text = mot
        row_cells[2].text = cap
    
    format_table(table)
    
    add_heading_custom(doc, "5.2 Threat Catalog and Risk Assessment", 2)
    
    # Threat table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Threat ID'
    hdr_cells[1].text = 'Threat Description'
    hdr_cells[2].text = 'Risk Level'
    hdr_cells[3].text = 'Primary Mitigation Controls'
    
    threats = [
        ("T1", "Firmware Injection – Malicious code inserted into update pipeline", "Critical", "ECDSA signing, HSM verification, Secure boot"),
        ("T2", "Man-in-the-Middle – Update packages intercepted and modified", "High", "TLS 1.3 mutual auth, Certificate pinning"),
        ("T3", "Replay Attack – Previously captured legitimate update retransmitted", "High", "Nonce-based freshness verification"),
        ("T4", "Rollback Attack – Forced reversion to firmware with known vulnerabilities", "High", "Monotonic version counter (HMS-stored)"),
        ("T5", "Supply Chain Compromise – Tier-1 supplier infrastructure compromised", "High", "SBOM validation, Supplier signing"),
        ("T6", "Key Compromise – Private code-signing keys obtained by attacker", "High", "HSM storage, MFA, Key rotation"),
        ("T7", "Denial of Service – Infrastructure flooded preventing legitimate updates", "High", "Rate limiting, CDN redundancy"),
        ("T8", "Rogue Update Server – Fraudulent server distributing malicious packages", "Medium", "Certificate pinning, Mutual TLS"),
        ("T9", "CAN Bus Injection – Exploitation of CAN bus authentication absence", "High", "Gateway ECU firewall, CAN anomaly detection"),
        ("T10", "Data Exfiltration – Extraction of firmware, keys, or personal data", "Medium", "TLS 1.3 encryption, Data minimization")
    ]
    
    for tid, desc, risk, controls in threats:
        row_cells = table.add_row().cells
        row_cells[0].text = tid
        row_cells[1].text = desc
        row_cells[2].text = risk
        row_cells[3].text = controls
    
    format_table(table)

def create_security_architecture(doc):
    """Create security architecture section"""
    add_heading_custom(doc, "6. Security Architecture and Technical Controls", 1)
    add_heading_custom(doc, "6.1 Defense-in-Depth Architecture", 2)
    
    add_paragraph_custom(doc, "The security architecture follows a defense-in-depth philosophy, implementing multiple overlapping security layers so that compromise of any single layer does not cause complete security failure (NIST, 2024). Controls are organized into four defensive layers.")
    
    # Architecture layers table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    layers = [
        ("Layer 1: Cloud-Side Security", "Secure build pipeline (CI/CD with integrity checks), ECDSA code signing (P-256 curve, HSM-backed), SBOM generation and validation, Access control (RBAC, MFA, separation of duties), Update campaign management (phased rollout)"),
        ("Layer 2: Transport Security", "TLS 1.3 with mutual authentication, Certificate pinning, Nonce-based freshness verification, Encrypted payload delivery via CDN"),
        ("Layer 3: Vehicle-Side Security", "HSM-backed ECDSA signature verification, Secure boot chain of trust, Monotonic version counter (anti-rollback, HSM-stored), A/B partition scheme (safe failover), Gateway ECU access control and network segmentation"),
        ("Layer 4: Monitoring and Response", "Lightweight ML-based IDS on update behavior, CAN bus anomaly detection, Cloud-side fleet telemetry monitoring, Automated rollback triggers, SIEM integration and SOC alerting")
    ]
    
    table.rows[0].cells[0].text = "Layer"
    table.rows[0].cells[1].text = "Security Controls"
    
    for i, (layer, controls) in enumerate(layers, 1):
        table.rows[i].cells[0].text = layer
        table.rows[i].cells[1].text = controls
    
    format_table(table)
    
    add_heading_custom(doc, "6.2 Cryptographic Controls", 2)
    add_paragraph_custom(doc, "Elliptic Curve Digital Signature Algorithm (ECDSA) with NIST P-256 curve provides 128-bit security with significantly faster verification than RSA, suitable for resource-constrained ECUs (Barker, 2020). The signing workflow includes: (1) firmware binary hashing using SHA-256; (2) hash signing using the OEM's private key stored in a FIPS 140-2 Level 3 certified HSM; (3) bundling signature, public key certificate, and metadata (version number, target ECU, timestamp, nonce) with the firmware; and (4) transferring the signed package to CDN for distribution. Upon receiving an update package, the TCU performs verification through the HSM using stored OEM public keys (NIST, 2018).")
    
    add_heading_custom(doc, "6.3 Control-to-Threat Mapping", 2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Threat'
    hdr_cells[1].text = 'Primary Controls'
    hdr_cells[2].text = 'Supporting Controls'
    
    mappings = [
        ("T1: Firmware Injection", "ECDSA signing, HSM verification, Secure boot", "SBOM validation, IDS monitoring"),
        ("T2: MITM Attack", "TLS 1.3 mutual auth, Certificate pinning", "Nonce freshness, IDS"),
        ("T3-T4: Replay/Rollback", "Nonce freshness, Monotonic version counter", "A/B partitioning, Secure boot"),
        ("T5: Supply Chain", "SBOM validation, Supplier signing", "Build pipeline integrity, RBAC"),
        ("T6: Key Compromise", "HSM storage, Separation of duties, MFA", "Key rotation, Audit logging"),
        ("T7-T10: Other Threats", "Rate limiting, Encryption, Data minimization", "Anomaly detection, Gateway controls")
    ]
    
    for threat, primary, supporting in mappings:
        row_cells = table.add_row().cells
        row_cells[0].text = threat
        row_cells[1].text = primary
        row_cells[2].text = supporting
    
    format_table(table)

def create_process_governance(doc):
    """Create process and governance section"""
    add_heading_custom(doc, "7. Process and Governance Model", 1)
    add_heading_custom(doc, "7.1 Cybersecurity Management System (CSMS)", 2)
    
    add_paragraph_custom(doc, "In compliance with UNECE WP.29 R155, the governance model is structured around a Cybersecurity Management System (CSMS) providing organizational accountability, defined processes, and continuous improvement for OTA security throughout the vehicle lifecycle (UNECE, 2020). The CSMS ensures cybersecurity is embedded as an ongoing operational discipline, not a one-time design activity.")
    
    add_heading_custom(doc, "7.2 Lifecycle Security Framework", 2)
    
    items = [
        "Design Phase: Threat Analysis and Risk Assessment (TARA) following ISO/SAE 21434 Clause 15; security requirements derived from TARA; secure coding standards (MISRA C/C++, CERT); Privacy Impact Assessment (PIA) for GDPR compliance (ISO/SAE, 2021; NIST, 2020).",
        
        "Production Phase: HSM-backed code signing; SBOM validation cross-referenced against vulnerability databases; security testing including positive functional and negative security tests; independent third-party penetration testing prior to deployment (ISO/SAE, 2021).",
        
        "Operation Phase: Phased rollout with canary group (1–5%), progressive expansion, and full fleet deployment after monitoring confirmation; continuous monitoring through the Vehicle Security Operations Center (VSOC); expedited patching for critical vulnerabilities within 72 hours of validation (UNECE, 2020).",
        
        "Decommissioning Phase: Secure key and certificate revocation; standards-compliant data erasure from vehicle and cloud systems; tamper-evident audit trail archival."
    ]
    add_bullet_list(doc, items)
    
    add_heading_custom(doc, "7.3 Incident Response Plan", 2)
    add_paragraph_custom(doc, "Following NIST SP 800-61 Rev. 2 (Cichonski et al., 2012), the incident response process includes: (1) Preparation through VSOC establishment and playbook development; (2) Detection and Analysis via IDS alerts, fleet anomalies, and threat intelligence; (3) Containment through update campaign halts and network blocking; (4) Eradication and Recovery via corrective firmware deployment; and (5) Post-Incident Activity including lessons-learned reviews and TARA updates. Regulatory notification occurs within mandated timelines (72 hours under GDPR; per WP.29 requirements).")

def create_business_considerations(doc):
    """Create business considerations section"""
    add_heading_custom(doc, "8. Business and Strategic Considerations", 1)
    
    add_heading_custom(doc, "8.1 Regulatory Compliance as Market Access Requirement", 2)
    add_paragraph_custom(doc, "Compliance with UNECE WP.29 R155 is mandatory for market access in the European Union, Japan, and South Korea. Non-compliance results in type approval denial, market exclusion, financial penalties, mandatory recalls, and strengthened liability exposure (UNECE, 2020). The proposed strategy explicitly supports R155 compliance by implementing all required CSMS elements.")
    
    add_heading_custom(doc, "8.2 Cost-Benefit Analysis", 2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Investment Category'
    hdr_cells[1].text = 'Investment Level'
    hdr_cells[2].text = 'Business Benefit'
    
    costs = [
        ("HSM Infrastructure (cloud + per-vehicle)", "High (capital + per-unit)", "Avoided recall costs ($100M+ per major incident) (Miller & Valasek, 2015)"),
        ("Secure Build Pipeline and CI/CD Security", "Moderate (one-time + maintenance)", "Regulatory compliance and market access"),
        ("VSOC Establishment and Operations", "High (ongoing)", "Brand trust and customer retention"),
        ("Third-Party Penetration Testing", "Moderate (annual)", "Reduced insurance premiums, Risk validation"),
        ("Training and Awareness Programs", "Low (ongoing)", "Reduced human-error incidents")
    ]
    
    for inv, cost, benefit in costs:
        row_cells = table.add_row().cells
        row_cells[0].text = inv
        row_cells[1].text = cost
        row_cells[2].text = benefit
    
    format_table(table)
    
    add_heading_custom(doc, "8.3 Strategic Value and Implementation Roadmap", 2)
    add_paragraph_custom(doc, "Secure OTA enables software-defined vehicle business models including feature-on-demand, subscription-based capabilities, and continuous vehicle improvement programs. Competitive differentiation grows as consumer cybersecurity awareness increases (McKinsey & Company, 2023). The phased implementation roadmap includes: Phase 1 (0–12 months) – HSM deployment, secure build pipeline, TLS 1.3 implementation, and CSMS documentation; Phase 2 (12–24 months) – ECDSA signing operational, A/B partitioning, phased rollout capability, and VSOC establishment; Phase 3 (24–36 months) – ML-based IDS, CAN anomaly detection, and SIEM integration; Phase 4 (36+ months) – Continuous improvement, post-quantum cryptography readiness, and V2X security planning.")

def create_validation(doc):
    """Create validation section"""
    add_heading_custom(doc, "9. Validation, Feasibility, and Limitations", 1)
    
    add_heading_custom(doc, "9.1 Validation Approach", 2)
    add_paragraph_custom(doc, "Standards alignment review confirms comprehensive coverage of ISO/SAE 21434, UNECE WP.29 R155, NIST SP 800-193, and NIST CSF 2.0. Threat-control coverage analysis demonstrates every identified threat is addressed by at least one primary and supporting control. Technology readiness assessment confirms all proposed technologies are production-deployed (TLS 1.3, ECDSA, HSMs, A/B partitioning, secure boot) or mature for near-term deployment (lightweight ML-based IDS, SBOM tooling) (Barker, 2020; Rescorla, 2018).")
    
    add_heading_custom(doc, "9.2 Feasibility Assessment", 2)
    items = [
        "Technical Feasibility: HIGH — All controls use established algorithms and commercially available hardware (Barker, 2020; NIST, 2018).",
        "Organizational Feasibility: MODERATE to HIGH — VSOC establishment and secure coding training align with WP.29 R155 transformation requirements (UNECE, 2020).",
        "Economic Feasibility: HIGH — Investment justified by risk mitigation and market access preservation; costs distributed over multi-year cycles.",
        "Regulatory Feasibility: HIGH — Strategy explicitly designed for compliance with applicable frameworks."
    ]
    add_bullet_list(doc, items)
    
    add_heading_custom(doc, "9.3 Limitations and Future Considerations", 2)
    add_paragraph_custom(doc, "Limitations include HSM dependency for older vehicles without hardware root-of-trust; CAN bus protocol limitations lacking native authentication and encryption; evolving threat landscape requiring post-quantum cryptography migration; supply chain depth challenges across multi-tier automotive suppliers; and residual insider threat risk despite access controls. Future considerations encompass post-quantum migration to CRYSTALS-Dilithium (NIST, 2022), Vehicle-to-Everything (V2X) security integration, AI/ML model update security, Zero Trust Architecture adoption, and international regulatory harmonization efforts.")

def create_conclusion(doc):
    """Create conclusion section"""
    add_heading_custom(doc, "10. Conclusion and Future Considerations", 1)
    
    paras = [
        "This report presented a comprehensive automobility cybersecurity strategy for securing Over-the-Air (OTA) updates, integrating technical controls, ethical considerations, governance processes, and business strategy into a cohesive, implementable plan. The core contributions include: (1) a rigorous threat model using STRIDE methodology identifying ten critical threats across the OTA ecosystem; (2) a defense-in-depth security architecture implementing four protective layers that collectively address all identified threats with multiple overlapping protections; (3) an ethical framework prioritizing human safety, privacy, fairness, and transparency; (4) a lifecycle governance model compliant with UNECE WP.29 R155 CSMS requirements; and (5) a business case demonstrating return on investment through regulatory compliance, avoided recall costs exceeding $150 million, competitive differentiation, and enablement of emerging software-defined vehicle revenue models (McKinsey & Company, 2023; Miller & Valasek, 2015).",
        
        "The strategy demonstrates that securing OTA updates is not merely a technical challenge but a multifaceted endeavor requiring the integration of engineering, ethics, governance, and business strategy. The automotive industry's transition to software-defined vehicles will succeed only if the update mechanisms enabling this transformation are trustworthy, resilient, and aligned with the safety expectations that society rightly demands of transportation systems (UNECE, 2020; ISO/SAE, 2021).",
        
        "Future work must address post-quantum cryptography migration as NIST finalizes standards, V2X security integration as vehicle connectivity expands, AI/ML model security for autonomous driving updates, Zero Trust Architecture principles, and industry-wide OTA security certification frameworks. Collaborative defense through expanded Auto-ISAC participation and information sharing will strengthen the sector's collective resilience against evolving threats. The stakes—human safety, privacy, and the future of mobility—justify sustained commitment to automotive cybersecurity excellence."
    ]
    
    for para in paras:
        add_paragraph_custom(doc, para)

def create_references(doc):
    """Create references section"""
    add_heading_custom(doc, "References", 1)
    
    references = [
        "Auto-ISAC. (2023). Automotive Information Sharing and Analysis Center: Best practices and threat intelligence sharing. https://automotiveisac.com/",
        
        "Barker, E. (2020). Recommendation for key management: Part 1 – General (NIST Special Publication 800-57 Part 1, Revision 5). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-57pt1r5",
        
        "Checkoway, S., McCoy, D., Kantor, B., Anderson, D., Shacham, H., Savage, S., Koscher, K., Czeskis, A., Roesner, F., & Kohno, T. (2011). Comprehensive experimental analyses of automotive attack surfaces. Proceedings of the 20th USENIX Security Symposium, 77–92.",
        
        "Cichonski, P., Millar, T., Grance, T., & Scarfone, K. (2012). Computer security incident handling guide (NIST Special Publication 800-61, Revision 2). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-61r2",
        
        "ENISA. (2019). Cybersecurity challenges in the uptake of artificial intelligence in autonomous driving. European Union Agency for Cybersecurity. https://www.enisa.europa.eu/publications/smart-cars",
        
        "Halder, S., Ozdenizci Kose, B., Edwards, B., & Guvenc, I. (2020). Secure OTA software updates in connected vehicles: A survey. IEEE Access, 8, 220946–220965. https://doi.org/10.1109/ACCESS.2020.3043413",
        
        "ISO/SAE. (2021). ISO/SAE 21434:2021 — Road vehicles — Cybersecurity engineering. International Organization for Standardization. https://www.iso.org/standard/70918.html",
        
        "Koscher, K., Czeskis, A., Roesner, F., Patel, S., Kohno, T., Checkoway, S., McCoy, D., Kantor, B., Anderson, D., Shacham, H., & Savage, S. (2010). Experimental security analysis of a modern automobile. Proceedings of the 2010 IEEE Symposium on Security and Privacy, 447–462. https://doi.org/10.1109/SP.2010.34",
        
        "Lim, H. S. M., Taeihagh, A., & Tan, A. (2022). Governing autonomous vehicles: Emerging responses for safety, liability, privacy, cybersecurity, and industry risks. Transport Reviews, 42(3), 370–395.",
        
        "McKinsey & Company. (2023). The software-defined vehicle: Mapping the automotive software and electronics landscape through 2030. https://www.mckinsey.com/industries/automotive-and-assembly/our-insights/mapping-the-automotive-software-and-electronics-landscape",
        
        "Miller, C., & Valasek, C. (2015). Remote exploitation of an unaltered passenger vehicle. Black Hat USA 2015.",
        
        "NIST. (2018). Platform firmware resiliency guidelines (Special Publication 800-193). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-193",
        
        "NIST. (2020). NIST Privacy Framework: A tool for improving privacy through enterprise risk management (Version 1.0). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.CSWP.01162020",
        
        "NIST. (2022). Post-quantum cryptography: Selected algorithms 2022. National Institute of Standards and Technology. https://csrc.nist.gov/projects/post-quantum-cryptography",
        
        "NIST. (2024). The NIST Cybersecurity Framework (CSF) 2.0 (NIST CSWP 29). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.CSWP.29",
        
        "NTIA. (2021). The minimum elements for a software bill of materials (SBOM). National Telecommunications and Information Administration, U.S. Department of Commerce. https://www.ntia.gov/report/2021/minimum-elements-software-bill-materials-sbom",
        
        "Rescorla, E. (2018). The Transport Layer Security (TLS) Protocol Version 1.3 (RFC 8446). Internet Engineering Task Force. https://doi.org/10.17487/RFC8446",
        
        "UNECE. (2020). UN Regulation No. 155 – Cybersecurity and cybersecurity management system. https://unece.org/transport/documents/2021/03/standards/un-regulation-no-155-cyber-security-and-cyber-security",
        
        "Upstream Security. (2023). Global automotive cybersecurity report. https://upstream.auto/"
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)

def main():
    # Create document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Configure Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Generate all sections
    create_title_page(doc)
    create_toc(doc)
    create_executive_summary(doc)
    doc.add_page_break()
    create_introduction(doc)
    doc.add_page_break()
    create_system_description(doc)
    doc.add_page_break()
    create_security_objectives(doc)
    doc.add_page_break()
    create_ethical_assessment(doc)
    doc.add_page_break()
    create_threat_modeling(doc)
    doc.add_page_break()
    create_security_architecture(doc)
    doc.add_page_break()
    create_process_governance(doc)
    doc.add_page_break()
    create_business_considerations(doc)
    doc.add_page_break()
    create_validation(doc)
    doc.add_page_break()
    create_conclusion(doc)
    doc.add_page_break()
    create_references(doc)
    
    # Save document
    output_path = '/home/siva/.openclaw/workspace/CYB400_OTA_Security_Report_CORRECTED.docx'
    doc.save(output_path)
    print(f"Corrected report generated successfully: {output_path}")

if __name__ == '__main__':
    main()
