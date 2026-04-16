#!/usr/bin/env python3
"""
Generate CYB400 Automobility Cybersecurity Strategy Final Report - Condensed Version
Topic: Securing Over-the-Air (OTA) Updates
Target: Max 25 pages with professional formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_paragraph_spacing(paragraph, space_before=0, space_after=6, line_spacing=1.15):
    """Set paragraph spacing"""
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing

def add_heading_custom(doc, text, level):
    """Add a heading with consistent formatting"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        else:
            run.font.size = Pt(12)
            run.font.bold = True
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, alignment=None, first_line_indent=0.3):
    """Add a paragraph with custom formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if alignment:
        p.alignment = alignment
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    set_paragraph_spacing(p)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet_list(doc, items, font_size=11):
    """Add a bulleted list"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)

def add_numbered_list(doc, items, font_size=11):
    """Add a numbered list"""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)

def set_cell_border(cell, **kwargs):
    """Set cell border properties"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_elm = OxmlElement(f'w:{edge}')
        edge_elm.set(qn('w:val'), 'single')
        edge_elm.set(qn('w:sz'), '4')
        edge_elm.set(qn('w:color'), '000000')
        tcBorders.append(edge_elm)
    tcPr.append(tcBorders)

def format_table(table):
    """Format table with consistent styling"""
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

def create_title_page(doc):
    """Create the title page"""
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Automobility Cybersecurity Strategy:\nSecuring Over-the-Air (OTA) Updates")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Final Project Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("Course: CYB400 Automobility Cybersecurity Strategy")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("Zekelman School of Business | St. Clair College")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    group = doc.add_paragraph()
    group.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = group.add_run("Group 8")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    
    members_text = "Siva Sai Akunuru | Manish Reddy Ennedla\nSri Venkata Naga Sai Chinmai Malladi | Lotachi Obi-Nwaigwe"
    members = doc.add_paragraph()
    members.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = members.add_run(members_text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    prof = doc.add_paragraph()
    prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prof.add_run("Professor: A. Sodiq Shofoluwe")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("Date: April 2026")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def create_toc(doc):
    """Create table of contents"""
    add_heading_custom(doc, "Table of Contents", 1)
    
    toc_items = [
        ("Executive Summary", 3),
        ("1. Introduction and Background", 5),
        ("2. System Description and Operating Context", 7),
        ("3. Security Objectives and Scope", 10),
        ("4. Ethical and Societal Impact Assessment", 12),
        ("5. Threat Modeling and Risk Analysis", 14),
        ("6. Security Architecture and Technical Controls", 17),
        ("7. Process and Governance Model", 20),
        ("8. Business and Strategic Considerations", 22),
        ("9. Validation, Feasibility, and Limitations", 24),
        ("10. Conclusion and Future Considerations", 25),
        ("References", 27),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops
        run = p.add_run(f"{item}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        
        # Add dots and page number
        tab_run = p.add_run("\t" * 8 + str(page))
        tab_run.font.name = 'Times New Roman'
        tab_run.font.size = Pt(11)
    
    doc.add_page_break()

def create_executive_summary(doc):
    """Create executive summary section"""
    add_heading_custom(doc, "Executive Summary", 1)
    
    paras = [
        "The automotive industry is undergoing a foundational transformation as vehicles evolve from mechanical systems to software-defined platforms. Over-the-Air (OTA) updates enable manufacturers to deploy firmware patches, security fixes, and feature enhancements remotely without dealership visits. Industry projections indicate that by 2030, over 80% of vehicle software will be distributed via OTA channels, making OTA security foundational to modern automobility (McKinsey & Company, 2023).",
        
        "This report presents a comprehensive cybersecurity strategy for securing OTA update systems in connected vehicles. The strategy addresses the complete ecosystem encompassing OEM cloud infrastructure, telematics control units (TCUs), electronic control units (ECUs), in-vehicle networks, and wireless communication channels. The security architecture employs defense-in-depth principles, integrating cryptographic authenticity verification, encrypted transmission protocols, secure boot validation, and anomaly detection mechanisms.",
        
        "Key security objectives include: (1) ensuring authenticity and integrity through ECDSA cryptographic signing with Hardware Security Module (HSM) backing; (2) protecting confidentiality via TLS 1.3 encryption; (3) preventing replay and rollback attacks through nonce-based freshness verification and monotonic version counters; (4) detecting anomalous behavior using lightweight ML-based intrusion detection; and (5) enabling safe recovery through automated rollback to known-good firmware.",
        
        "The strategy demonstrates alignment with ISO/SAE 21434 for automotive cybersecurity engineering, UNECE WP.29 R155 for vehicle cybersecurity regulation, and NIST SP 800-193 for platform firmware resiliency. The ethical framework prioritizes human safety and privacy through privacy-by-design principles, user consent mechanisms, and fail-safe protections. The business case demonstrates that security investment is justified by regulatory compliance, avoided recall costs exceeding $150 million for major incidents like the 2015 Jeep Cherokee recall, competitive differentiation, and enablement of emerging software-defined vehicle revenue models.",
        
        "This strategy provides a defensible, implementable framework for securing OTA updates that balances technical rigor, ethical responsibility, and business practicality in the rapidly evolving landscape of connected and autonomous vehicles."
    ]
    
    for para in paras:
        add_paragraph_custom(doc, para, first_line_indent=0.3)

def create_introduction(doc):
    """Create introduction section"""
    add_heading_custom(doc, "1. Introduction and Background", 1)
    add_heading_custom(doc, "1.1 The Software-Defined Vehicle Revolution", 2)
    
    add_paragraph_custom(doc, "The automotive industry is undergoing foundational transformation. Modern vehicles are increasingly defined by software governing ADAS, powertrain management, infotainment, and autonomous driving functions. Contemporary premium vehicles contain over 100 million lines of code across 100+ ECUs, with complexity accelerating (McKinsey & Company, 2023). This evolution enables the software-defined vehicle (SDV) paradigm, where functionality can be continuously enhanced post-manufacturing through OTA updates.")
    
    add_heading_custom(doc, "1.2 The OTA Update Paradigm", 2)
    add_paragraph_custom(doc, "OTA updates encompass Firmware-Over-the-Air (FOTA) for ECU firmware and Software-Over-the-Air (SOTA) for application-layer software. Updates leverage cellular (4G/5G) and Wi-Fi to transmit packages from OEM cloud infrastructure to vehicle fleets. The process involves: (1) package creation and cryptographic signing at the OEM backend; (2) cloud distribution; (3) encrypted wireless transmission; (4) TCU reception and verification; (5) ECU distribution via CAN bus; and (6) post-installation validation (ISO/SAE, 2021).")
    
    add_heading_custom(doc, "1.3 The Cybersecurity Imperative", 2)
    add_paragraph_custom(doc, "OTA capability creates critical attack surfaces exploitable at scale. In 2015, researchers Miller and Valasek demonstrated remote exploitation of a Jeep Cherokee, gaining wireless control over steering, braking, and transmission through connectivity vulnerabilities—leading to recall of 1.4 million vehicles (Miller & Valasek, 2015). Automotive cybersecurity incidents continue increasing, with remote attacks constituting the majority and OTA vulnerabilities representing growing concern (Upstream Security, 2023).")
    
    add_heading_custom(doc, "1.4 Regulatory and Standards Landscape", 2)
    add_paragraph_custom(doc, "Regulators have established comprehensive frameworks: UNECE WP.29 R155 mandates certified Cybersecurity Management Systems (CSMS) (UNECE, 2020); ISO/SAE 21434 provides automotive cybersecurity engineering standards (ISO/SAE, 2021); NIST SP 800-193 addresses firmware resiliency (NIST, 2018); and the NIST Cybersecurity Framework provides organizational structure (NIST, 2024).")

def create_system_description(doc):
    """Create system description section"""
    add_heading_custom(doc, "2. System Description and Operating Context", 1)
    add_heading_custom(doc, "2.1 OTA Ecosystem Architecture", 2)
    
    add_paragraph_custom(doc, "The OTA update ecosystem spans cloud infrastructure, wireless networks, and in-vehicle computing platforms (ISO/SAE, 2021). Figure 1 illustrates the high-level architecture comprising cloud-side components, communication channels, and vehicle-side systems.")
    
    # Architecture diagram as text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Figure 1: OTA Update Ecosystem Architecture")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    diagram = """
    ┌─────────────────────────────────────────────────────────────┐
    │                    OEM CLOUD INFRASTRUCTURE                  │
    │  ┌────────────┐  ┌──────────────┐  ┌──────────────┐         │
    │  │Build & Sign│→ │Update Mgmt   │→ │MQTT/HTTPS    │         │
    │  │Server (HSM)│  │Platform      │  │Broker        │         │
    │  └────────────┘  └──────────────┘  └──────────────┘         │
    └─────────────────────────────────┬───────────────────────────┘
                                      │ TLS 1.3
                                      ▼
    ┌─────────────────────────────────────────────────────────────┐
    │                     VEHICLE PLATFORM                         │
    │  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐     │
    │  │   TCU    │→ │ Gateway  │→ │ Target   │  │   HSM    │     │
    │  │(Telematic)│  │   ECU    │  │   ECUs   │  │(Security)│     │
    │  └──────────┘  └──────────┘  └──────────┘  └──────────┘     │
    │                    CAN Bus / Automotive Ethernet             │
    └─────────────────────────────────────────────────────────────┘
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(diagram)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    add_heading_custom(doc, "2.2 Cloud-Side Components", 2)
    items = [
        "Build and Signing Server: Secure environment for firmware compilation and ECDSA cryptographic signing using HSM-backed private keys (NIST, 2018).",
        "Update Management Platform: Orchestrates update campaigns, phased rollouts, and deployment monitoring (ISO/SAE, 2021).",
        "MQTT/HTTPS Broker: Facilitates lightweight publish-subscribe communication between cloud and vehicles (Halder et al., 2020).",
        "Content Delivery Network (CDN): Geographically distributed update package distribution."
    ]
    add_bullet_list(doc, items)
    
    add_heading_custom(doc, "2.3 Vehicle-Side Components", 2)
    items = [
        "Telematics Control Unit (TCU): Primary connectivity gateway receiving updates and coordinating ECU distribution (Upstream Security, 2023).",
        "Gateway ECU: Firewall and traffic controller enforcing network segmentation between safety-critical and non-critical domains (Checkoway et al., 2011).",
        "Target ECUs: Individual controllers receiving firmware updates for ADAS, powertrain, body, and infotainment systems.",
        "Hardware Security Module (HSM): Tamper-resistant hardware providing secure key storage and signature verification (NIST, 2018)."
    ]
    add_bullet_list(doc, items)

def create_security_objectives(doc):
    """Create security objectives section"""
    add_heading_custom(doc, "3. Security Objectives and Scope", 1)
    add_heading_custom(doc, "3.1 Security Objectives", 2)
    
    objectives = [
        "Authenticity and Integrity Assurance: Ensure every OTA update is authentic and unmodified through ECDSA cryptographic signing with HSM-backed keys (NIST, 2018).",
        "Confidentiality of Update Payloads: Protect update content during transmission via TLS 1.3 encryption and at-rest encryption.",
        "Replay and Rollback Prevention: Prevent adversaries from retransmitting captured updates or forcing firmware reversion through nonce-based freshness verification and monotonic version counters stored in HSM.",
        "Anomaly Detection and Monitoring: Detect anomalous update behavior using lightweight ML-based intrusion detection systems at vehicle and fleet levels.",
        "Safe Recovery and Resilience: Enable automated rollback to known-good firmware via A/B partition schemes and fail-safe modes maintaining essential vehicle functionality (NIST, 2018)."
    ]
    add_numbered_list(doc, objectives)
    
    add_heading_custom(doc, "3.2 Scope and Exclusions", 2)
    
    # Scope table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'In-Scope Elements'
    hdr_cells[1].text = 'Exclusions'
    
    scope_data = [
        ("OEM cloud infrastructure (build servers, signing infrastructure, update platform)", "Physical hardware root-of-trust at silicon level"),
        ("Cellular (4G/5G) and Wi-Fi communication channels", "Non-OTA mechanisms (USB, dealership diagnostic tools)"),
        ("TCU, Gateway ECU, target ECUs, HSM", "Vehicle-to-Everything (V2X) communications security"),
        ("CAN bus and automotive Ethernet networks", "Autonomous driving algorithm functional safety validation"),
        ("OEMs, Tier-1/Tier-2 suppliers, vehicle owners, regulators", "Legacy vehicles without connectivity or HSM capability")
    ]
    
    for in_scope, exclusion in scope_data:
        row_cells = table.add_row().cells
        row_cells[0].text = in_scope
        row_cells[1].text = exclusion
    
    format_table(table)

def create_ethical_assessment(doc):
    """Create ethical assessment section"""
    add_heading_custom(doc, "4. Ethical and Societal Impact Assessment", 1)
    
    add_paragraph_custom(doc, "OTA security carries profound ethical implications given the safety-critical nature of automotive systems operating in shared public spaces (UNECE, 2020). This section examines ethical dimensions and strategy responses grounded in NIST Privacy Framework principles (NIST, 2020).")
    
    add_heading_custom(doc, "4.1 Safety and Duty of Care", 2)
    add_paragraph_custom(doc, "OEMs bear duty of care to ensure vehicles do not pose unreasonable risks. When OTA updates modify safety-critical systems (ADAS, braking, steering), compromised updates could cause fatalities (Miller & Valasek, 2015). Strategy response includes: cryptographic signature verification ensuring only authorized firmware installation; dual-bank A/B partitioning guaranteeing known-good firmware fallback; fail-safe modes maintaining essential functions during failures; and phased rollout policies deploying to canary groups before wider release (ISO/SAE, 2021).")
    
    add_heading_custom(doc, "4.2 Privacy and Data Protection", 2)
    add_paragraph_custom(doc, "OTA processes involve vehicle data including VINs, firmware versions, and location. Strategy response includes: data minimization collecting only necessary information with telemetry anonymization; informed consent with explicit owner notification; TLS 1.3 encryption for all data in transit; retention limits minimizing personal data storage; and compliance with GDPR, CCPA, and NIST Privacy Framework (NIST, 2020).")
    
    add_heading_custom(doc, "4.3 Fairness, Transparency, and Societal Trust", 2)
    add_paragraph_custom(doc, "Security updates must be distributed equitably without discrimination based on vehicle tier or geography. The strategy mandates security-critical patches for all affected vehicles regardless of model or subscription status. Transparency requirements include clear release notes, publicly accessible cybersecurity disclosure policies, and timely incident notification. These measures maintain societal trust in connected mobility essential for technology adoption (Upstream Security, 2023).")

def create_threat_modeling(doc):
    """Create threat modeling section"""
    add_heading_custom(doc, "5. Threat Modeling and Risk Analysis", 1)
    add_heading_custom(doc, "5.1 Methodology and Adversary Profiles", 2)
    
    add_paragraph_custom(doc, "Threat modeling employs the STRIDE framework (Spoofing, Tampering, Repudiation, Information Disclosure, Denial of Service, Elevation of Privilege) adapted for automotive contexts per ISO/SAE 21434 TARA methodology (ISO/SAE, 2021; ENISA, 2019).")
    
    # Adversary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Adversary Type'
    hdr_cells[1].text = 'Motivation'
    hdr_cells[2].text = 'Capability'
    
    adversaries = [
        ("Script Kiddies", "Curiosity, notoriety", "Low; public tools"),
        ("Organized Criminals", "Financial gain", "Moderate-High; custom malware, ransomware"),
        ("Nation-State Actors", "Espionage, sabotage", "Very High; APT, zero-days (ENISA, 2019)"),
        ("Malicious Insiders", "Financial gain, revenge", "High; privileged access"),
        ("Security Researchers", "Public safety", "High; deep expertise (Miller & Valasek, 2015)")
    ]
    
    for adv, mot, cap in adversaries:
        row_cells = table.add_row().cells
        row_cells[0].text = adv
        row_cells[1].text = mot
        row_cells[2].text = cap
    
    format_table(table)
    
    add_heading_custom(doc, "5.2 Threat Catalog and Risk Assessment", 2)
    
    # Threat table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Threat'
    hdr_cells[1].text = 'STRIDE'
    hdr_cells[2].text = 'Risk Level'
    hdr_cells[3].text = 'Primary Controls'
    
    threats = [
        ("T1: Firmware Injection", "Tampering, EoP", "Critical", "ECDSA signing, HSM verification, Secure boot"),
        ("T2: MITM Attack", "Tampering, ID", "High", "TLS 1.3 mutual auth, Certificate pinning"),
        ("T3: Replay Attack", "Spoofing, Tampering", "High", "Nonce-based freshness verification"),
        ("T4: Rollback Attack", "Tampering", "High", "Monotonic version counter (HSM-stored)"),
        ("T5: Supply Chain Compromise", "Tampering, EoP", "High", "SBOM validation, Supplier signing"),
        ("T6: Key Compromise", "Spoofing", "High", "HSM storage, MFA, Key rotation"),
        ("T7: Denial of Service", "DoS", "High", "Rate limiting, CDN redundancy"),
        ("T8: Rogue Update Server", "Spoofing", "Medium", "Certificate pinning, Mutual TLS"),
        ("T9: CAN Bus Injection", "Tampering, EoP", "High", "Gateway firewall, CAN anomaly detection"),
        ("T10: Data Exfiltration", "ID", "Medium", "TLS 1.3, Data minimization")
    ]
    
    for threat, stride, risk, controls in threats:
        row_cells = table.add_row().cells
        row_cells[0].text = threat
        row_cells[1].text = stride
        row_cells[2].text = risk
        row_cells[3].text = controls
    
    format_table(table)

def create_security_architecture(doc):
    """Create security architecture section"""
    add_heading_custom(doc, "6. Security Architecture and Technical Controls", 1)
    add_heading_custom(doc, "6.1 Defense-in-Depth Architecture", 2)
    
    add_paragraph_custom(doc, "The security architecture follows defense-in-depth principles with four overlapping layers ensuring compromise of any single layer does not cause complete failure (NIST, 2024).")
    
    # Architecture description as table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    layers = [
        ("Layer 1: Cloud-Side", "Secure build pipeline, ECDSA code signing (P-256, HSM-backed), SBOM generation, RBAC with MFA, phased rollout management"),
        ("Layer 2: Transport", "TLS 1.3 with mutual authentication, Certificate pinning, Nonce-based freshness, Encrypted CDN delivery"),
        ("Layer 3: Vehicle-Side", "HSM-backed ECDSA verification, Secure boot chain of trust, Monotonic version counter, A/B partition scheme, Gateway ECU firewall"),
        ("Layer 4: Monitoring", "ML-based IDS on update behavior, CAN bus anomaly detection, Cloud fleet telemetry monitoring, Automated rollback triggers, SIEM integration")
    ]
    
    for i, (layer, controls) in enumerate(layers):
        table.rows[i].cells[0].text = layer
        table.rows[i].cells[1].text = controls
    
    format_table(table)
    
    add_heading_custom(doc, "6.2 Cryptographic Controls", 2)
    add_paragraph_custom(doc, "ECDSA with NIST P-256 curve provides 128-bit security with faster verification than RSA, suitable for resource-constrained ECUs (Barker, 2020). The signing workflow: (1) SHA-256 hash of firmware; (2) HSM-backed private key signing; (3) Bundle signature, certificate, metadata (version, timestamp, nonce); (4) CDN distribution. Verification at TCU uses HSM-stored OEM public keys (NIST, 2018).")
    
    add_heading_custom(doc, "6.3 Control-to-Threat Mapping", 2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Threat'
    hdr_cells[1].text = 'Primary Controls'
    hdr_cells[2].text = 'Supporting Controls'
    
    mappings = [
        ("T1: Firmware Injection", "ECDSA signing, HSM verification, Secure boot", "SBOM validation, IDS monitoring"),
        ("T2: MITM Attack", "TLS 1.3 mutual auth, Certificate pinning", "Nonce freshness, IDS"),
        ("T3-T4: Replay/Rollback", "Nonce freshness, Monotonic version counter", "A/B partitioning, Secure boot"),
        ("T5: Supply Chain", "SBOM validation, Supplier signing", "Build pipeline integrity, RBAC"),
        ("T6: Key Compromise", "HSM storage, Separation of duties, MFA", "Key rotation, Audit logging"),
        ("T7-T10: Other", "Rate limiting, Encryption, Data minimization", "Anomaly detection, Gateway controls")
    ]
    
    for threat, primary, supporting in mappings:
        row_cells = table.add_row().cells
        row_cells[0].text = threat
        row_cells[1].text = primary
        row_cells[2].text = supporting
    
    format_table(table)

def create_process_governance(doc):
    """Create process and governance section"""
    add_heading_custom(doc, "7. Process and Governance Model", 1)
    add_heading_custom(doc, "7.1 Cybersecurity Management System (CSMS)", 2)
    
    add_paragraph_custom(doc, "Compliance with UNECE WP.29 R155 requires a CSMS providing organizational accountability and continuous improvement throughout the vehicle lifecycle (UNECE, 2020). The lifecycle framework spans Design (TARA, secure coding standards, privacy impact assessment), Production (code signing, SBOM validation, penetration testing), Operation (phased rollout, continuous monitoring, expedited patching), and Decommissioning (key revocation, data erasure).")
    
    add_heading_custom(doc, "7.2 Incident Response Plan", 2)
    add_paragraph_custom(doc, "Following NIST SP 800-61 Rev. 2 (Cichonski et al., 2012), the incident response process includes: (1) Preparation through VSOC establishment and playbook development; (2) Detection and Analysis via IDS alerts, fleet anomalies, and threat intelligence; (3) Containment through update campaign halts and network blocking; (4) Eradication and Recovery via corrective firmware deployment; and (5) Post-Incident Activity including lessons-learned reviews and TARA updates. Regulatory notification occurs within mandated timelines (72 hours under GDPR; per WP.29 requirements).")
    
    add_heading_custom(doc, "7.3 Stakeholder Roles (RACI)", 2)
    
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    
    activities = ["Threat Modeling", "Code Signing", "Deployment", "Incident Response", "Compliance Audits"]
    oem = ["A/R", "A", "A/R", "A/R", "A"]
    tier1 = ["C", "R", "I", "C", "R"]
    dealer = ["—", "—", "C", "I", "C"]
    regulator = ["I", "—", "I", "I", "R"]
    
    # Header row
    table.rows[0].cells[0].text = "Activity"
    table.rows[0].cells[1].text = "OEM"
    table.rows[0].cells[2].text = "Tier-1"
    table.rows[0].cells[3].text = "Dealer"
    table.rows[0].cells[4].text = "Regulator"
    
    for i, activity in enumerate(activities):
        table.rows[i+1].cells[0].text = activity
        table.rows[i+1].cells[1].text = oem[i]
        table.rows[i+1].cells[2].text = tier1[i]
        table.rows[i+1].cells[3].text = dealer[i]
        table.rows[i+1].cells[4].text = regulator[i]
    
    format_table(table)
    
    p = doc.add_paragraph()
    p.add_run("R = Responsible, A = Accountable, C = Consulted, I = Informed").font.size = Pt(9)

def create_business_considerations(doc):
    """Create business considerations section"""
    add_heading_custom(doc, "8. Business and Strategic Considerations", 1)
    
    add_heading_custom(doc, "8.1 Regulatory Compliance and Market Access", 2)
    add_paragraph_custom(doc, "UNECE WP.29 R155 compliance is mandatory for market access in EU, Japan, and South Korea. Non-compliance results in type approval denial, financial penalties, and strengthened liability exposure (UNECE, 2020).")
    
    add_heading_custom(doc, "8.2 Cost-Benefit Analysis", 2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Investment Area'
    hdr_cells[1].text = 'Cost Level'
    hdr_cells[2].text = 'Benefit'
    
    costs = [
        ("HSM Infrastructure", "High", "Avoided recalls ($100M+ per major incident) (Miller & Valasek, 2015)"),
        ("Secure Build Pipeline", "Moderate", "Regulatory compliance, Market access"),
        ("VSOC Operations", "High (ongoing)", "Brand trust, Customer retention"),
        ("Penetration Testing", "Moderate (annual)", "Reduced insurance premiums, Risk validation"),
        ("Training/Programs", "Low", "Reduced human-error incidents")
    ]
    
    for inv, cost, benefit in costs:
        row_cells = table.add_row().cells
        row_cells[0].text = inv
        row_cells[1].text = cost
        row_cells[2].text = benefit
    
    format_table(table)
    
    add_heading_custom(doc, "8.3 Strategic Value and Roadmap", 2)
    add_paragraph_custom(doc, "Secure OTA enables software-defined vehicle business models including feature-on-demand and subscription services. Competitive differentiation grows as consumer cybersecurity awareness increases (McKinsey & Company, 2023). The phased implementation roadmap includes: Phase 1 (0-12 months): HSM deployment, secure build pipeline, TLS 1.3; Phase 2 (12-24 months): ECDSA signing operational, A/B partitioning, VSOC establishment; Phase 3 (24-36 months): ML-based IDS, CAN anomaly detection, SIEM integration; Phase 4 (36+ months): Continuous improvement, post-quantum readiness.")

def create_validation(doc):
    """Create validation section"""
    add_heading_custom(doc, "9. Validation, Feasibility, and Limitations", 1)
    
    add_heading_custom(doc, "9.1 Validation Approach", 2)
    add_paragraph_custom(doc, "Standards alignment review confirms comprehensive coverage of ISO/SAE 21434, UNECE WP.29 R155, NIST SP 800-193, and NIST CSF 2.0. Threat-control coverage analysis demonstrates every identified threat is addressed by at least one primary and supporting control. Proof-of-concept prototyping using open-source tools validated ECDSA P-256 signing (sub-second verification), MQTT-over-TLS 1.3 communication, nonce-based freshness rejection of replays, and ML-based CAN anomaly detection.")
    
    add_heading_custom(doc, "9.2 Feasibility Assessment", 2)
    items = [
        "Technical Feasibility: HIGH — All controls use established algorithms and commercially available hardware (Barker, 2020; NIST, 2018).",
        "Organizational Feasibility: MODERATE-HIGH — VSOC establishment and secure coding training align with WP.29 R155 transformation (UNECE, 2020).",
        "Economic Feasibility: HIGH — Investment justified by risk mitigation and market access; costs distributed over multi-year cycles.",
        "Regulatory Feasibility: HIGH — Strategy explicitly designed for compliance with applicable frameworks."
    ]
    add_bullet_list(doc, items)
    
    add_heading_custom(doc, "9.3 Limitations and Future Considerations", 2)
    add_paragraph_custom(doc, "Limitations include: HSM dependency for older vehicles; CAN bus protocol limitations lacking native authentication; evolving threat landscape requiring post-quantum cryptography migration (CRYSTALS-Dilithium); supply chain depth challenges; and insider threat residual risk. Future considerations encompass post-quantum migration (NIST, 2022), V2X security integration, AI/ML model security, zero trust architecture exploration, and international regulatory harmonization.")

def create_conclusion(doc):
    """Create conclusion section"""
    add_heading_custom(doc, "10. Conclusion and Future Considerations", 1)
    
    paras = [
        "This report presented a comprehensive automobility cybersecurity strategy for securing OTA updates, integrating technical controls, ethical considerations, governance processes, and business strategy. The core contributions include: (1) A rigorous threat model using STRIDE methodology identifying ten critical threats; (2) A defense-in-depth security architecture with four protective layers addressing all identified threats; (3) An ethical framework prioritizing safety, privacy, fairness, and transparency; (4) A lifecycle governance model compliant with UNECE WP.29 R155 CSMS requirements; and (5) A business case demonstrating ROI through regulatory compliance, avoided recall costs exceeding $150M, and competitive differentiation.",
        
        "The strategy demonstrates that securing OTA updates requires integration of engineering, ethics, governance, and business strategy. The automotive industry's transition to software-defined vehicles will succeed only if update mechanisms are trustworthy, resilient, and aligned with societal safety expectations (UNECE, 2020; ISO/SAE, 2021).",
        
        "Future work must address post-quantum cryptography migration, V2X security integration, AI/ML model security, zero trust architecture adoption, and industry-wide certification frameworks. Collaborative defense through Auto-ISAC participation and expanded information sharing will strengthen sector-wide resilience. The stakes—human safety, privacy, and the future of mobility—justify sustained commitment to automotive cybersecurity excellence."
    ]
    
    for para in paras:
        add_paragraph_custom(doc, para, first_line_indent=0.3)

def create_references(doc):
    """Create references section"""
    add_heading_custom(doc, "References", 1)
    
    references = [
        "Auto-ISAC. (2023). Automotive Information Sharing and Analysis Center: Best practices. https://automotiveisac.com/",
        "Barker, E. (2020). Recommendation for key management: Part 1 – General (NIST SP 800-57). NIST.",
        "Checkoway, S., et al. (2011). Comprehensive experimental analyses of automotive attack surfaces. USENIX Security.",
        "Cichonski, P., et al. (2012). Computer security incident handling guide (NIST SP 800-61 Rev. 2). NIST.",
        "ENISA. (2019). Cybersecurity challenges in AI for autonomous driving. EU Agency for Cybersecurity.",
        "Halder, S., et al. (2020). Secure OTA software updates in connected vehicles: A survey. IEEE Access.",
        "ISO/SAE. (2021). ISO/SAE 21434: Road vehicles — Cybersecurity engineering.",
        "Koscher, K., et al. (2010). Experimental security analysis of a modern automobile. IEEE S&P.",
        "McKinsey & Company. (2023). The software-defined vehicle through 2030.",
        "Miller, C., & Valasek, C. (2015). Remote exploitation of an unaltered passenger vehicle. Black Hat USA.",
        "NIST. (2018). Platform firmware resiliency guidelines (SP 800-193).",
        "NIST. (2020). NIST Privacy Framework: A tool for improving privacy through enterprise risk management.",
        "NIST. (2022). Post-quantum cryptography: Selected algorithms.",
        "NIST. (2024). The NIST Cybersecurity Framework (CSF) 2.0.",
        "NTIA. (2021). The minimum elements for a software bill of materials (SBOM).",
        "Rescorla, E. (2018). The Transport Layer Security Protocol Version 1.3 (RFC 8446). IETF.",
        "UNECE. (2020). UN Regulation No. 155 – Cybersecurity and cybersecurity management system.",
        "Upstream Security. (2023). Global automotive cybersecurity report."
    ]
    
    for ref in references:
        p = doc.add_paragraph(style='List Number')
        p.add_run(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)

def main():
    # Create document
    doc = Document()
    
    # Set page margins (standard academic format)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
    
    # Configure Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Generate all sections
    create_title_page(doc)
    create_toc(doc)
    create_executive_summary(doc)
    doc.add_page_break()
    create_introduction(doc)
    doc.add_page_break()
    create_system_description(doc)
    doc.add_page_break()
    create_security_objectives(doc)
    doc.add_page_break()
    create_ethical_assessment(doc)
    doc.add_page_break()
    create_threat_modeling(doc)
    doc.add_page_break()
    create_security_architecture(doc)
    doc.add_page_break()
    create_process_governance(doc)
    doc.add_page_break()
    create_business_considerations(doc)
    doc.add_page_break()
    create_validation(doc)
    doc.add_page_break()
    create_conclusion(doc)
    doc.add_page_break()
    create_references(doc)
    
    # Save document
    output_path = '/home/siva/.openclaw/workspace/CYB400_OTA_Security_Report_Final.docx'
    doc.save(output_path)
    print(f"Report generated successfully: {output_path}")

if __name__ == '__main__':
    main()
