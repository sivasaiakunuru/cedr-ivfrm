#!/usr/bin/env python3
"""
Generate CYB400 Professional Final Report - PhD Level
Topic: Securing Over-the-Air (OTA) Updates in Automotive Cybersecurity
Features: Professional TOC, embedded diagrams, academic rigor
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement, qn

def add_toc(doc):
    """Add a proper Table of Contents field"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # Create TOC field
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    # Add placeholder text
    placeholder = paragraph.add_run("[Table of Contents - Right-click and select 'Update Field' to generate]")
    placeholder.font.name = 'Times New Roman'
    placeholder.font.size = Pt(11)
    placeholder.italic = True
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    return paragraph

def set_cell_border(cell):
    """Set cell border properties"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_elm = OxmlElement(f'w:{edge}')
        edge_elm.set(qn('w:val'), 'single')
        edge_elm.set(qn('w:sz'), '6')
        edge_elm.set(qn('w:color'), '000000')
        tcBorders.append(edge_elm)
    tcPr.append(tcBorders)

def format_table(table, header_color=(0, 51, 102)):
    """Format table with consistent styling and header shading"""
    table.style = 'Table Grid'
    
    # Format header row
    header_cells = table.rows[0].cells
    for cell in header_cells:
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set header background color
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '%02x%02x%02x' % header_color)
        tcPr.append(shd)
    
    # Format data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

def add_heading_custom(doc, text, level):
    """Add a heading with professional formatting"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        else:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, first_line_indent=0.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a paragraph with academic formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Inches(first_line_indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = alignment
    return p

def add_figure_caption(doc, caption):
    """Add a figure caption"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {caption}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    run.font.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

def add_architecture_diagram(doc):
    """Add the OTA architecture diagram"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ASCII diagram with better formatting
    diagram = """
    ┌─────────────────────────────────────────────────────────────────────┐
    │                     OEM CLOUD INFRASTRUCTURE                         │
    │  ┌─────────────────┐  ┌──────────────────┐  ┌─────────────────┐    │
    │  │   BUILD &       │  │   UPDATE         │  │   MQTT/HTTPS    │    │
    │  │   SIGN SERVER   │→ │   MANAGEMENT     │→ │   BROKER        │    │
    │  │   (HSM-Backed)  │  │   PLATFORM       │  │                 │    │
    │  └─────────────────┘  └──────────────────┘  └─────────────────┘    │
    │           │                    │                     │              │
    │           └────────────────────┴─────────────────────┘              │
    └─────────────────────────────────┬───────────────────────────────────┘
                                      │
                                      ▼ TLS 1.3 (Encrypted)
    ┌─────────────────────────────────────────────────────────────────────┐
    │                      VEHICLE PLATFORM                                │
    │                                                                     │
    │   ┌──────────────┐    ┌──────────────┐    ┌──────────────┐        │
    │   │     TCU      │→   │   GATEWAY    │→   │   TARGET     │        │
    │   │  (Primary    │    │     ECU      │    │    ECUs      │        │
    │   │  Connectivity│    │  (Firewall   │    │ (ADAS, Body, │        │
    │   │   Gateway)   │    │   & Router)  │    │  Powertrain) │        │
    │   └──────────────┘    └──────────────┘    └──────────────┘        │
    │          │                   │                   │                 │
    │          └───────────────────┴───────────────────┘                 │
    │                              │                                     │
    │                              ▼                                     │
    │                   ┌──────────────────────┐                        │
    │                   │  CAN Bus / Ethernet  │                        │
    │                   └──────────────────────┘                        │
    │                                                                     │
    │   ┌─────────────────────────────────────────────────────────┐     │
    │   │              HARDWARE SECURITY MODULE (HSM)              │     │
    │   │         (Secure Key Storage & Verification)              │     │
    │   └─────────────────────────────────────────────────────────┘     │
    └─────────────────────────────────────────────────────────────────────┘
    """
    run = p.add_run(diagram)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    add_figure_caption(doc, "1: OTA Update Ecosystem Architecture")

def add_attack_tree_diagram(doc):
    """Add attack tree diagram"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    diagram = """
                         Firmware Injection Attack
                                  │
            ┌─────────────────────┼─────────────────────┐
            │                     │                     │
    ┌───────▼───────┐    ┌────────▼────────┐   ┌───────▼───────┐
    │  Compromise   │    │  Intercept &    │   │  Exploit TCU  │
    │  Build Server │    │  Modify in      │   │  Vulnerability│
    └───────┬───────┘    │  Transit        │   └───────┬───────┘
            │            └────────┬────────┘           │
    ┌───────┴───────┐             │            ┌───────┴───────┐
    │               │             │            │               │
┌───▼───┐     ┌────▼────┐  ┌─────▼─────┐  ┌───▼───┐     ┌────▼────┐
│Insider│     │Supply   │  │   MITM    │  │Buffer │     │  Auth   │
│Access │     │Chain    │  │  Attack   │  │Overflow│    │ Bypass  │
└───────┘     └─────────┘  └───────────┘  └───────┘     └─────────┘
    """
    run = p.add_run(diagram)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    add_figure_caption(doc, "2: Attack Tree – Firmware Injection (T1)")

def add_defense_layers_diagram(doc):
    """Add defense-in-depth layers diagram"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    diagram = """
    ╔═══════════════════════════════════════════════════════════════════╗
    ║  LAYER 4: MONITORING & RESPONSE                                   ║
    ║  ┌──────────────────────────────────────────────────────────┐     ║
    ║  │  • ML-Based IDS    • CAN Anomaly Detection               │     ║
    ║  │  • Fleet Telemetry • SIEM Integration                    │     ║
    ║  │  • Automated Rollback Triggers                           │     ║
    ║  └──────────────────────────────────────────────────────────┘     ║
    ╠═══════════════════════════════════════════════════════════════════╣
    ║  LAYER 3: VEHICLE-SIDE SECURITY                                   ║
    ║  ┌──────────────────────────────────────────────────────────┐     ║
    ║  │  • HSM Verification  • Secure Boot Chain                 │     ║
    ║  │  • Version Counters  • A/B Partitioning                  │     ║
    ║  │  • Gateway Firewall  • Network Segmentation              │     ║
    ║  └──────────────────────────────────────────────────────────┘     ║
    ╠═══════════════════════════════════════════════════════════════════╣
    ║  LAYER 2: TRANSPORT SECURITY                                      ║
    ║  ┌──────────────────────────────────────────────────────────┐     ║
    ║  │  • TLS 1.3 (Mutual Auth)  • Certificate Pinning          │     ║
    ║  │  • Freshness Verification • Encrypted CDN Delivery       │     ║
    ║  └──────────────────────────────────────────────────────────┘     ║
    ╠═══════════════════════════════════════════════════════════════════╣
    ║  LAYER 1: CLOUD-SIDE SECURITY                                     ║
    ║  ┌──────────────────────────────────────────────────────────┐     ║
    ║  │  • Secure Build Pipeline  • ECDSA Signing (HSM-backed)   │     ║
    ║  │  • SBOM Validation        • RBAC & MFA                   │     ║
    ║  │  • Phased Rollout Management                             │     ║
    ║  └──────────────────────────────────────────────────────────┘     ║
    ╚═══════════════════════════════════════════════════════════════════╝
    """
    run = p.add_run(diagram)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    add_figure_caption(doc, "3: Defense-in-Depth Security Architecture")

def add_lifecycle_diagram(doc):
    """Add security lifecycle diagram"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    diagram = """
    ┌──────────────┐      ┌──────────────┐      ┌──────────────┐      ┌──────────────┐
    │   DESIGN     │  →   │  PRODUCTION  │  →   │  OPERATION   │  →   │ DECOMMISION  │
    │              │      │              │      │              │      │              │
    │ • TARA       │      │ • Code       │      │ • Phased     │      │ • Key        │
    │ • Security   │      │   Signing    │      │   Rollout    │      │   Revocation │
    │   Req'ments  │      │ • SBOM       │      │ • Monitoring │      │ • Data       │
    │ • Secure     │      │   Validation │      │ • Patching   │      │   Erasure    │
    │   Coding     │      │ • Pen Test   │      │ • Incident   │      │ • Audit      │
    │              │      │              │      │   Response   │      │   Archive    │
    └──────────────┘      └──────────────┘      └──────────────┘      └──────────────┘
    """
    run = p.add_run(diagram)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    add_figure_caption(doc, "4: OTA Security Lifecycle Framework")

def create_title_page(doc):
    """Create the professional title page"""
    for _ in range(8):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AUTOMOBILITY CYBERSECURITY STRATEGY:")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Comprehensive Framework for Securing")
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("Over-the-Air (OTA) Update Systems")
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("CYB400 – Automobility Cybersecurity Strategy")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("Zekelman School of Business")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    school2 = doc.add_paragraph()
    school2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school2.add_run("St. Clair College")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    group = doc.add_paragraph()
    group.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = group.add_run("Group 8")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    members = doc.add_paragraph()
    members.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = members.add_run("Siva Sai Akunuru | Manish Reddy Ennedla\nSri Venkata Naga Sai Chinmai Malladi | Lotachi Obi-Nwaigwe")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    prof = doc.add_paragraph()
    prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prof.add_run("Professor: A. Sodiq Shofoluwe")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("April 2026")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def create_executive_summary(doc):
    """Create executive summary with academic tone"""
    add_heading_custom(doc, "Executive Summary", 1)
    
    paras = [
        "The paradigm shift toward software-defined vehicles (SDVs) has fundamentally transformed the automotive industry's operational landscape. Over-the-Air (OTA) update mechanisms have emerged as the principal vector through which Original Equipment Manufacturers (OEMs) deliver firmware patches, security fixes, feature enhancements, and performance optimizations to connected vehicle fleets at scale. Industry analyses project that by 2030, in excess of 80 percent of all automotive software modifications will be disseminated via OTA channels, rendering the security of these update systems foundational to modern automobility (McKinsey & Company, 2023).",
        
        "This report presents a comprehensive cybersecurity strategy for securing OTA update ecosystems within connected vehicles. The proposed framework addresses the complete technological stack, encompassing OEM cloud infrastructure, telematics control units (TCUs), electronic control units (ECUs), in-vehicle networks, and wireless communication channels. The security architecture employs defense-in-depth principles, integrating cryptographic authenticity verification, encrypted transmission protocols, secure boot validation mechanisms, and anomaly detection systems.",
        
        "The strategy is organized around five foundational security objectives: (1) ensuring authenticity and integrity of OTA packages through Elliptic Curve Digital Signature Algorithm (ECDSA) cryptographic signing with Hardware Security Module (HSM) backing; (2) protecting confidentiality via Transport Layer Security (TLS) 1.3 encryption; (3) preventing replay and rollback attacks through nonce-based freshness verification and monotonic version counters; (4) detecting anomalous behavior using lightweight machine learning-based intrusion detection; and (5) enabling safe recovery through automated rollback to known-good firmware states.",
        
        "The proposed framework demonstrates comprehensive alignment with ISO/SAE 21434 for automotive cybersecurity engineering, UNECE WP.29 Regulation No. 155 for vehicle cybersecurity management systems, and NIST Special Publication 800-193 for platform firmware resiliency. The ethical framework prioritizes human safety and privacy through privacy-by-design principles, informed consent mechanisms, and fail-safe architectural protections. The accompanying business case demonstrates that cybersecurity investment yields positive returns through regulatory compliance, avoidance of recall costs exceeding USD $150 million (as exemplified by the 2015 Jeep Cherokee incident), competitive differentiation, and enablement of emerging software-defined vehicle revenue models.",
        
        "This strategy provides a defensible, implementable framework for securing OTA updates that appropriately balances technical rigor, ethical responsibility, and business pragmatism within the rapidly evolving landscape of connected and autonomous vehicles."
    ]
    
    for para in paras:
        add_paragraph_custom(doc, para)

def create_introduction(doc):
    """Create introduction with academic rigor"""
    add_heading_custom(doc, "1. Introduction and Background", 1)
    
    add_heading_custom(doc, "1.1 The Software-Defined Vehicle Revolution", 2)
    
    add_paragraph_custom(doc, "The automotive industry is currently experiencing a paradigm shift of unprecedented magnitude. Modern vehicles are increasingly characterized by their software architecture rather than purely mechanical attributes, with software governing Advanced Driver-Assistance Systems (ADAS), powertrain management, infotainment systems, and autonomous driving functionalities. Contemporary premium-class vehicles incorporate in excess of 100 million lines of source code distributed across more than 100 Electronic Control Units (ECUs), with software complexity metrics demonstrating consistent annual growth rates exceeding 25 percent (McKinsey & Company, 2023). This evolution has precipitated the emergence of the software-defined vehicle paradigm, wherein vehicle functionality, performance characteristics, and safety capabilities can be continuously modified, enhanced, and repaired throughout the operational lifecycle via Over-the-Air (OTA) update mechanisms.")
    
    add_heading_custom(doc, "1.2 The OTA Update Mechanism", 2)
    
    add_paragraph_custom(doc, "OTA update mechanisms within the automotive domain encompass two primary classifications: Firmware-Over-the-Air (FOTA) updates targeting low-level ECU firmware and safety-critical software components, and Software-Over-the-Air (SOTA) updates addressing application-layer software including infotainment systems, navigation platforms, and user interface components. Both categories leverage wireless communication infrastructures—primarily cellular networks (4G LTE and 5G) and IEEE 802.11 Wi-Fi protocols—to transmit cryptographically packaged updates from OEM cloud infrastructure to individual vehicles or fleet populations. The canonical OTA process encompasses: (1) package creation and cryptographic signing within the OEM secure build environment; (2) distribution through cloud-based content delivery infrastructure; (3) transmission over encrypted wireless channels; (4) reception and multi-stage verification at the vehicle TCU; (5) ECU-specific distribution via in-vehicle networks including Controller Area Network (CAN) bus and Automotive Ethernet; and (6) post-installation cryptographic validation and operational confirmation (ISO/SAE, 2021).")
    
    add_heading_custom(doc, "1.3 The Cybersecurity Imperative", 2)
    
    add_paragraph_custom(doc, "While OTA technology delivers substantial operational and economic benefits, it simultaneously creates a critical attack surface exploitable at unprecedented scale. The remote, wireless nature of OTA mechanisms implies that a successful compromise could potentially affect millions of vehicles simultaneously across geographically dispersed populations. In 2015, security researchers Miller and Valasek demonstrated the remote exploitation of an unaltered Jeep Cherokee, successfully gaining wireless control over steering, braking, and transmission subsystems through vulnerabilities in the vehicle's connectivity infrastructure (Miller & Valasek, 2015). This proof-of-concept demonstration precipitated the recall of 1.4 million vehicles and catalyzed industry-wide recognition of automotive cybersecurity as a safety-critical concern. Subsequent analyses indicate that remote attack vectors now constitute the majority of reported automotive cybersecurity incidents, with OTA-related vulnerabilities representing a rapidly growing proportion of disclosed Common Vulnerabilities and Exposures (CVEs) (Upstream Security, 2023).")
    
    add_heading_custom(doc, "1.4 Regulatory and Standards Framework", 2)
    
    add_paragraph_custom(doc, "In recognition of escalating threat landscapes, regulatory bodies and standards organizations have established comprehensive frameworks addressing automotive cybersecurity. The United Nations Economic Commission for Europe (UNECE) WP.29 World Forum for Harmonization of Vehicle Regulations has promulgated Regulation No. 155, which mandates that vehicle manufacturers implement a certified Cybersecurity Management System (CSMS) and demonstrate systematic cybersecurity risk management throughout the vehicle lifecycle (UNECE, 2020). The International Organization for Standardization (ISO) and Society of Automotive Engineers (SAE) Joint Working Group has developed ISO/SAE 21434, providing a comprehensive engineering standard for automotive cybersecurity encompassing threat analysis, risk assessment, security requirement specification, design, implementation, verification, and validation activities (ISO/SAE, 2021). The National Institute of Standards and Technology (NIST) has published Special Publication 800-193, which offers authoritative guidance on platform firmware resiliency, specifically addressing the protection, detection, and recovery of firmware from unauthorized modification (NIST, 2018).")

def create_system_description(doc):
    """Create system description with diagram"""
    add_heading_custom(doc, "2. System Description and Operating Context", 1)
    
    add_heading_custom(doc, "2.1 Architectural Overview", 2)
    
    add_paragraph_custom(doc, "The OTA update ecosystem constitutes a distributed, multi-layered cyber-physical system spanning cloud computing infrastructure, wireless telecommunications networks, and embedded automotive computing platforms. A comprehensive understanding of this architecture is prerequisite for systematic vulnerability identification and appropriate security control design (ISO/SAE, 2021). Figure 1 illustrates the high-level ecosystem architecture.")
    
    # Add architecture diagram
    add_architecture_diagram(doc)
    
    add_heading_custom(doc, "2.2 Cloud-Side Infrastructure", 2)
    
    items = [
        "Build and Signing Server: The OEM's secure build environment wherein firmware packages undergo compilation, quality assurance testing, and cryptographic signing utilizing private keys maintained within FIPS 140-2 Level 3 certified Hardware Security Modules (HSMs). The signing process constitutes a critical trust anchor, ensuring that only authorized, cryptographically verified software enters the distribution pipeline (NIST, 2018).",
        
        "Update Management Platform: The orchestration layer responsible for comprehensive update campaign management, including targeting specific vehicle models, firmware versions, geographic regions, and rollout cohorts. This platform manages phased deployment strategies, monitors deployment metrics and success rates, and triggers automated rollback procedures upon detection of anomalous installation patterns (ISO/SAE, 2021).",
        
        "Message Broker: The communication middleware facilitating bidirectional messaging between cloud infrastructure and vehicle fleets. Message Queuing Telemetry Transport (MQTT) protocol implementations are widely deployed in automotive OTA systems due to the protocol's lightweight publish-subscribe architecture, minimized bandwidth requirements, and suitability for resource-constrained embedded environments (Halder et al., 2020).",
        
        "Content Delivery Network (CDN): Geographically distributed infrastructure for update package dissemination, ensuring low-latency, high-availability delivery to vehicle populations across diverse geographic regions."
    ]
    
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
    
    add_heading_custom(doc, "2.3 Vehicle-Side Components", 2)
    
    items = [
        "Telematics Control Unit (TCU): The vehicle's primary external connectivity gateway, responsible for receiving encrypted update packages from cloud infrastructure, performing cryptographic verification, and coordinating distribution to target ECUs. The TCU constitutes the primary security boundary between external untrusted networks and the internal vehicle network (Upstream Security, 2023).",
        
        "Gateway ECU: Acts as an internal firewall, traffic controller, and protocol translator for the in-vehicle network. The Gateway enforces network segmentation policies and access control rules, isolating safety-critical domains (powertrain, braking, steering) from non-critical domains (infotainment, telematics) to prevent lateral movement by adversaries (Checkoway et al., 2011).",
        
        "Target Electronic Control Units (ECUs): The individual embedded controllers receiving and installing firmware updates. Target ECUs encompass safety-critical systems (ADAS, anti-lock braking, electronic stability control), body control modules, powertrain controllers, battery management systems (in electric vehicles), and infotainment processors.",
        
        "Hardware Security Module (HSM): A tamper-resistant cryptographic processor providing secure key storage, hardware-accelerated cryptographic operations, and secure boot enforcement. The HSM serves as the vehicle's hardware root of trust for OTA update verification and constitutes the foundational security element for the entire vehicle cybersecurity architecture (NIST, 2018)."
    ]
    
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)

def create_security_objectives(doc):
    """Create security objectives"""
    add_heading_custom(doc, "3. Security Objectives and Scope", 1)
    
    add_heading_custom(doc, "3.1 Security Objectives", 2)
    
    objectives = [
        "Authenticity and Integrity Assurance: Ensure that every OTA update package received by a vehicle is authentic (originating from a legitimate OEM) and maintains cryptographic integrity (unmodified during transit and storage). This objective is achieved through ECDSA P-256 cryptographic code signing with private keys maintained in HSMs at the cloud side and corresponding public key verification within vehicle-side HSMs (NIST, 2018).",
        
        "Confidentiality of Update Payloads: Protect the informational content of OTA update packages from unauthorized disclosure during transmission and storage. Confidentiality protection is implemented via TLS 1.3 encrypted transport channels providing forward secrecy, coupled with at-rest encryption for packages stored within cloud infrastructure.",
        
        "Replay and Rollback Attack Prevention: Prevent adversaries from retransmitting previously captured legitimate update packages (replay attacks) or forcing the installation of outdated firmware versions containing known security vulnerabilities (rollback attacks). Mitigation employs cryptographically secure nonce-based freshness verification and monotonic version counters maintained within tamper-resistant HSM storage.",
        
        "Anomaly Detection and Security Monitoring: Detect anomalous or malicious behavior associated with OTA processes at both the individual vehicle level and the fleet-wide aggregate level. Detection capabilities leverage lightweight machine learning-based intrusion detection systems optimized for resource-constrained automotive computing environments.",
        
        "Safe Recovery and Operational Resilience: Ensure vehicles can safely recover from failed, interrupted, or compromised update operations through automated rollback mechanisms to known-good firmware states. Resilience is achieved via A/B partition schemes, watchdog timers, and fail-safe operational modes that maintain essential vehicle functionality during fault conditions (NIST, 2018)."
    ]
    
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f"Objective {i}: {obj}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.25)
    
    add_heading_custom(doc, "3.2 Scope Definition and Exclusions", 2)
    
    # Scope table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'In-Scope Elements'
    hdr_cells[1].text = 'Out-of-Scope Elements'
    
    scope_data = [
        ("OEM cloud infrastructure (secure build servers, HSM-backed signing infrastructure, update orchestration platforms)", "Physical hardware root-of-trust implementation at the semiconductor fabrication and chip design level"),
        ("Cellular (4G/5G) and Wi-Fi communication channels between cloud and vehicle", "Non-OTA update mechanisms including USB-based flashing, dealership diagnostic tool programming, and JTAG interfaces"),
        ("Vehicle-side components: TCU, Gateway ECU, target ECUs, HSM, in-vehicle networks", "Vehicle-to-Everything (V2X) communications security including V2V and V2I protocols"),
        ("CAN bus, CAN-FD, and Automotive Ethernet in-vehicle networks", "Functional safety validation of autonomous driving decision algorithms (addressed by ISO 26262)"),
        ("Stakeholders: OEMs, Tier-1/Tier-2 suppliers, vehicle owners, regulatory authorities", "Legacy vehicles manufactured prior to widespread OTA capability adoption")
    ]
    
    for in_scope, exclusion in scope_data:
        row_cells = table.add_row().cells
        row_cells[0].text = in_scope
        row_cells[1].text = exclusion
    
    format_table(table)

def create_threat_modeling(doc):
    """Create threat modeling with diagrams"""
    add_heading_custom(doc, "5. Threat Modeling and Risk Analysis", 1)
    
    add_heading_custom(doc, "5.1 Threat Modeling Methodology", 2)
    
    add_paragraph_custom(doc, "The threat modeling process employed in this analysis utilizes the STRIDE framework (Spoofing, Tampering, Repudiation, Information Disclosure, Denial of Service, Elevation of Privilege), adapted specifically for the automotive OTA context and supplemented by the European Union Agency for Cybersecurity (ENISA) threat taxonomy for smart vehicles (ENISA, 2019). This methodology aligns with the Threat Analysis and Risk Assessment (TARA) process specified in ISO/SAE 21434, Clause 15, which constitutes the industry-standard approach for automotive cybersecurity risk assessment (ISO/SAE, 2021).")
    
    add_heading_custom(doc, "5.2 Adversary Taxonomy", 2)
    
    # Adversary table
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Adversary Classification'
    hdr_cells[1].text = 'Primary Motivation'
    hdr_cells[2].text = 'Technical Capability'
    hdr_cells[3].text = 'Threat Sophistication'
    
    adversaries = [
        ("Opportunistic Attackers (Script Kiddies)", "Notoriety, curiosity, recreational hacking", "Low; reliance on publicly available exploit tools and automated attack frameworks", "Low"),
        ("Organized Cybercriminal Groups", "Financial gain through ransomware, data theft, fraud", "Moderate to High; development of custom malware, social engineering capabilities", "Moderate to High"),
        ("Nation-State Actors and APT Groups", "Espionage, infrastructure sabotage, geopolitical objectives", "Very High; zero-day vulnerability discovery, supply chain infiltration, advanced persistent threat capabilities (ENISA, 2019)", "Very High"),
        ("Malicious Insider Threats", "Financial gain, revenge, ideological motivations", "High; privileged access to internal systems, knowledge of operational procedures", "High"),
        ("Security Researchers (White Hat)", "Vulnerability discovery, public safety improvement, academic recognition", "High to Very High; deep technical expertise, advanced reverse engineering capabilities", "High")
    ]
    
    for adv, mot, cap, soph in adversaries:
        row_cells = table.add_row().cells
        row_cells[0].text = adv
        row_cells[1].text = mot
        row_cells[2].text = cap
        row_cells[3].text = soph
    
    format_table(table)
    
    add_heading_custom(doc, "5.3 Threat Catalog and Risk Assessment", 2)
    
    add_paragraph_custom(doc, "Systematic threat identification yielded ten critical threat categories targeting the OTA ecosystem. Table 2 presents the comprehensive threat catalog with corresponding risk assessments and primary mitigation controls.")
    
    # Threat table
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Threat ID'
    hdr_cells[1].text = 'Threat Description'
    hdr_cells[2].text = 'STRIDE Category'
    hdr_cells[3].text = 'Risk Level'
    hdr_cells[4].text = 'Primary Mitigation Controls'
    
    threats = [
        ("T1", "Firmware Injection: Malicious firmware insertion into the update pipeline", "Tampering, Elevation of Privilege", "Critical", "ECDSA P-256 signing, HSM-backed verification, Secure Boot chain"),
        ("T2", "Man-in-the-Middle: Interception and modification of update packages in transit", "Tampering, Information Disclosure", "High", "TLS 1.3 mutual authentication, Certificate pinning"),
        ("T3", "Replay Attack: Retransmission of previously captured legitimate updates", "Spoofing, Tampering", "High", "Cryptographic nonce freshness verification"),
        ("T4", "Rollback Attack: Forced installation of outdated, vulnerable firmware", "Tampering", "High", "Monotonic version counters in HSM-protected storage"),
        ("T5", "Supply Chain Compromise: Infiltration of Tier-1 supplier build systems", "Tampering, Elevation of Privilege", "High", "SBOM validation, Multi-party code signing"),
        ("T6", "Code Signing Key Compromise: Theft or unauthorized use of private signing keys", "Spoofing", "High", "HSM key storage, Multi-factor authentication, Separation of duties"),
        ("T7", "Denial of Service: Infrastructure saturation preventing legitimate updates", "Denial of Service", "High", "Rate limiting, CDN redundancy, DDoS mitigation"),
        ("T8", "Rogue Update Server: Fraudulent infrastructure distributing malicious packages", "Spoofing", "Medium", "Certificate pinning, Mutual TLS authentication"),
        ("T9", "In-Vehicle Network Injection: CAN bus exploitation post-update", "Tampering, Elevation of Privilege", "High", "Gateway ECU firewall, CAN anomaly detection"),
        ("T10", "Exfiltration: Unauthorized extraction of firmware, keys, or telemetry data", "Information Disclosure", "Medium", "TLS 1.3 encryption, Data minimization principles")
    ]
    
    for tid, desc, stride, risk, controls in threats:
        row_cells = table.add_row().cells
        row_cells[0].text = tid
        row_cells[1].text = desc
        row_cells[2].text = stride
        row_cells[3].text = risk
        row_cells[4].text = controls
    
    format_table(table)
    
    add_heading_custom(doc, "5.4 Attack Tree Analysis", 2)
    
    add_paragraph_custom(doc, "Attack tree modeling provides a systematic methodology for decomposing high-level adversary objectives into constituent attack paths, enabling comprehensive defense design. Figure 2 presents the attack tree for Threat T1 (Firmware Injection), which represents the most critical threat vector due to its potential for direct compromise of vehicle safety systems.")
    
    # Add attack tree diagram
    add_attack_tree_diagram(doc)

def create_security_architecture(doc):
    """Create security architecture section"""
    add_heading_custom(doc, "6. Security Architecture and Technical Controls", 1)
    
    add_heading_custom(doc, "6.1 Defense-in-Depth Architecture", 2)
    
    add_paragraph_custom(doc, "The proposed security architecture adheres to defense-in-depth principles, implementing multiple overlapping protective layers such that compromise of any individual layer does not result in complete security failure. This layered approach ensures defense resilience against sophisticated, multi-stage attacks (NIST, 2024). Figure 3 illustrates the four-layer defense architecture.")
    
    # Add defense layers diagram
    add_defense_layers_diagram(doc)
    
    add_heading_custom(doc, "6.2 Cryptographic Security Controls", 2)
    
    add_paragraph_custom(doc, "The Elliptic Curve Digital Signature Algorithm (ECDSA) with NIST P-256 curve parameters provides 128-bit security equivalence with significantly superior computational efficiency compared to RSA cryptosystems. A 256-bit ECDSA key offers security comparable to a 3072-bit RSA key while reducing signature verification time by approximately 60 percent—a critical consideration for resource-constrained automotive ECUs (Barker, 2020). The cryptographic workflow encompasses:")
    
    crypto_steps = [
        "Firmware binary processing through SHA-256 cryptographic hash function",
        "Hash signing utilizing the OEM's private ECDSA key maintained within FIPS 140-2 Level 3 certified HSM",
        "Bundle construction incorporating signature, X.509 public key certificate, and metadata (semantic version number, target ECU identifier, Unix timestamp, cryptographic nonce)",
        "Secure transfer of signed packages to geographically distributed CDN infrastructure",
        "Vehicle-side verification using HSM-stored OEM public keys with hardware-accelerated ECDSA verification"
    ]
    
    for step in crypto_steps:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(step)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.5)
    
    add_heading_custom(doc, "6.3 Threat-to-Control Mapping", 2)
    
    add_paragraph_custom(doc, "Table 3 presents the comprehensive mapping between identified threats and corresponding security controls, demonstrating complete threat coverage through primary and supporting control implementations.")
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Threat Classification'
    hdr_cells[1].text = 'Primary Security Controls'
    hdr_cells[2].text = 'Supporting Security Controls'
    
    mappings = [
        ("T1: Firmware Injection", "ECDSA P-256 signing, HSM-backed verification, Secure Boot chain validation", "SBOM validation, Build pipeline integrity monitoring, Runtime IDS"),
        ("T2: MITM Attack", "TLS 1.3 with mutual authentication, X.509 certificate pinning", "Nonce-based freshness verification, Network anomaly detection"),
        ("T3-T4: Replay/Rollback", "Cryptographic nonce freshness, Monotonic HSM-backed version counters", "A/B partition rollback capability, Secure Boot enforcement"),
        ("T5: Supply Chain", "Software Bill of Materials validation, Multi-party supplier signing", "Build environment integrity verification, RBAC with separation of duties"),
        ("T6: Key Management", "FIPS 140-2 Level 3 HSM storage, Multi-factor authentication, Separation of duties", "Regular key rotation, Comprehensive audit logging, Air-gapped backup"),
        ("T7-T10: Infrastructure", "Rate limiting and throttling, End-to-end encryption, Data minimization", "DDoS mitigation, Gateway access control, Behavioral anomaly detection")
    ]
    
    for threat, primary, supporting in mappings:
        row_cells = table.add_row().cells
        row_cells[0].text = threat
        row_cells[1].text = primary
        row_cells[2].text = supporting
    
    format_table(table)

def create_process_governance(doc):
    """Create process and governance section"""
    add_heading_custom(doc, "7. Process and Governance Model", 1)
    
    add_heading_custom(doc, "7.1 Cybersecurity Management System (CSMS)", 2)
    
    add_paragraph_custom(doc, "In compliance with UNECE WP.29 Regulation No. 155, the governance framework implements a certified Cybersecurity Management System (CSMS) providing organizational accountability, documented processes, and continuous improvement mechanisms for OTA security throughout the vehicle lifecycle (UNECE, 2020). The CSMS ensures cybersecurity is institutionalized as an ongoing operational discipline rather than treated as a discrete design-phase activity.")
    
    add_heading_custom(doc, "7.2 Security Lifecycle Framework", 2)
    
    add_paragraph_custom(doc, "Figure 4 illustrates the comprehensive security lifecycle framework spanning vehicle development, production, operation, and decommissioning phases.")
    
    # Add lifecycle diagram
    add_lifecycle_diagram(doc)
    
    items = [
        "Design Phase: Formal Threat Analysis and Risk Assessment (TARA) following ISO/SAE 21434 Clause 15; security requirements derivation and traceability through implementation; adoption of MISRA C/C++ and CERT secure coding standards; Privacy Impact Assessment (PIA) for regulatory compliance (ISO/SAE, 2021; NIST, 2020).",
        
        "Production Phase: All firmware packages undergo HSM-backed ECDSA signing; Software Bill of Materials generation and validation against National Vulnerability Database (NVD); comprehensive security testing including positive functional validation and negative security testing (malformed inputs, signature bypass attempts); independent third-party penetration testing prior to production authorization (ISO/SAE, 2021).",
        
        "Operation Phase: Phased rollout strategy beginning with canary deployment (1–5 percent of fleet), progressive expansion upon validation, and full fleet deployment following monitoring confirmation; continuous security monitoring through the Vehicle Security Operations Center (VSOC); expedited patching procedures for critical vulnerabilities targeting 72-hour deployment timelines (UNECE, 2020).",
        
        "Decommissioning Phase: Systematic revocation of cryptographic keys and certificates; cryptographic erasure of sensitive data from vehicle systems and cloud databases; tamper-evident archival of audit logs for regulatory retention periods."
    ]
    
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.25)
    
    add_heading_custom(doc, "7.3 Incident Response Framework", 2)
    
    add_paragraph_custom(doc, "The incident response framework adheres to the structured methodology specified in NIST Special Publication 800-61 Revision 2 (Cichonski et al., 2012), adapted for the automotive OTA context and integrated with UNECE WP.29 R155 reporting requirements. The response lifecycle encompasses: (1) Preparation through VSOC establishment, playbook development, and tabletop exercises; (2) Detection and Analysis utilizing IDS alerts, fleet telemetry anomalies, and threat intelligence feeds; (3) Containment via update campaign suspension, network-level isolation, and affected vehicle identification; (4) Eradication and Recovery through corrective firmware development, testing, and expedited deployment; (5) Post-Incident Activity including root cause analysis, TARA updates, and information sharing through Auto-ISAC. Regulatory notification obligations require disclosure within 72 hours for GDPR personal data breaches and per WP.29 timelines for vehicle safety incidents.")

def create_business_considerations(doc):
    """Create business considerations"""
    add_heading_custom(doc, "8. Business and Strategic Considerations", 1)
    
    add_heading_custom(doc, "8.1 Regulatory Compliance Imperative", 2)
    
    add_paragraph_custom(doc, "UNECE WP.29 Regulation No. 155 establishes mandatory type approval requirements for cybersecurity management systems. Non-compliance carries severe consequences including: type approval denial preventing market entry; exclusion from European Union, Japanese, and South Korean markets; substantial financial penalties; mandatory safety recalls; and significantly strengthened plaintiff claims in product liability litigation (UNECE, 2020). The proposed strategy explicitly addresses all CSMS requirements, ensuring regulatory compliance as a prerequisite for market access.")
    
    add_heading_custom(doc, "8.2 Economic Analysis", 2)
    
    add_paragraph_custom(doc, "Table 4 presents the comprehensive cost-benefit analysis for OTA cybersecurity investment.")
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Investment Category'
    hdr_cells[1].text = 'Investment Magnitude'
    hdr_cells[2].text = 'Quantified Benefit'
    
    costs = [
        ("Hardware Security Module Infrastructure (cloud + per-vehicle)", "High (capital expenditure + per-unit cost)", "Avoided recall costs: USD $100–500M per major incident; 2015 Jeep Cherokee recall exceeded $150M (Miller & Valasek, 2015)"),
        ("Secure Build Pipeline and CI/CD Security Integration", "Moderate (initial + ongoing maintenance)", "Regulatory compliance enabling market access in WP.29 jurisdictions"),
        ("Vehicle Security Operations Center (VSOC) Establishment", "High (personnel + infrastructure)", "Brand protection, customer retention, insurance premium reduction"),
        ("Third-Party Security Assessment and Penetration Testing", "Moderate (annual recurring)", "Independent validation, reduced cyber insurance premiums"),
        ("Personnel Training and Security Awareness Programs", "Low to Moderate (ongoing)", "Reduced human error incidents, improved incident response capability")
    ]
    
    for inv, cost, benefit in costs:
        row_cells = table.add_row().cells
        row_cells[0].text = inv
        row_cells[1].text = cost
        row_cells[2].text = benefit
    
    format_table(table)
    
    add_heading_custom(doc, "8.3 Strategic Value Proposition", 2)
    
    add_paragraph_custom(doc, "Secure OTA capability constitutes a foundational enabler for emerging software-defined vehicle business models, including feature-on-demand offerings, subscription-based capability upgrades, and continuous vehicle improvement programs. Competitive differentiation increasingly derives from demonstrable cybersecurity posture as consumer awareness of automotive security risks increases (McKinsey & Company, 2023). The phased implementation roadmap encompasses: Phase 1 (Months 0–12): HSM infrastructure deployment, secure build pipeline implementation, TLS 1.3 rollout, and CSMS documentation; Phase 2 (Months 12–24): ECDSA signing operationalization, A/B partitioning implementation, VSOC establishment; Phase 3 (Months 24–36): Machine learning-based IDS deployment, CAN anomaly detection, SIEM integration; Phase 4 (Months 36+): Continuous improvement, post-quantum cryptography preparation, and V2X security integration.")

def create_validation(doc):
    """Create validation section"""
    add_heading_custom(doc, "9. Validation, Feasibility, and Limitations", 1)
    
    add_heading_custom(doc, "9.1 Validation Methodology", 2)
    
    add_paragraph_custom(doc, "Comprehensive validation of the proposed strategy encompasses multiple verification dimensions. Standards alignment review confirms comprehensive coverage of ISO/SAE 21434 (Clause 8–15), UNECE WP.29 R155 Annex 5 (CSMS requirements), NIST SP 800-193 (firmware resiliency guidelines), and NIST CSF 2.0 (Identify, Protect, Detect, Respond, Recover functions). Threat-control coverage analysis demonstrates that each of the ten identified threats is addressed by at least one primary control and multiple supporting controls, ensuring defense-in-depth resilience. Technology readiness assessment confirms that all proposed controls utilize production-deployed technologies (TLS 1.3, ECDSA P-256, HSMs, secure boot) or technologies mature for near-term deployment (lightweight ML-based IDS, SBOM tooling) (Barker, 2020; Rescorla, 2018).")
    
    add_heading_custom(doc, "9.2 Feasibility Assessment", 2)
    
    items = [
        "Technical Feasibility: HIGH — All cryptographic algorithms, protocols, and hardware components are commercially available, extensively peer-reviewed, and deployed in production environments (Barker, 2020; NIST, 2018).",
        
        "Organizational Feasibility: MODERATE to HIGH — VSOC establishment and security culture transformation align with industry-wide transitions mandated by WP.29 R155. Organizational change management is required but supported by regulatory imperative (UNECE, 2020).",
        
        "Economic Feasibility: HIGH — Net present value analysis demonstrates positive returns when accounting for avoided recall costs, regulatory compliance value, insurance premium reductions, and enablement of new revenue streams. Implementation costs are distributed across multi-year deployment cycles.",
        
        "Regulatory Feasibility: HIGH — The strategy is explicitly architected to satisfy UNECE WP.29 R155, ISO/SAE 21434, GDPR, and applicable NIST framework requirements."
    ]
    
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.25)
    
    add_heading_custom(doc, "9.3 Limitations and Future Research", 2)
    
    add_paragraph_custom(doc, "Acknowledged limitations include: (1) Hardware Security Module dependency, wherein older vehicles lacking HSM capability cannot achieve equivalent security assurance; (2) Controller Area Network (CAN) protocol limitations, as the CAN bus lacks native authentication or encryption mechanisms, requiring compensating gateway controls; (3) post-quantum cryptography considerations, as the emergence of cryptographically relevant quantum computers threatens current ECDSA algorithms within the 15–20 year vehicle operational lifetime, necessitating eventual migration to quantum-resistant algorithms such as CRYSTALS-Dilithium (NIST, 2022); (4) supply chain visibility constraints, as complete security assurance across multi-tier automotive supply chains remains an industry-wide challenge; and (5) residual insider threat risk, as determined adversaries with authorized access may circumvent controls over extended periods despite monitoring systems. Future research directions encompass Vehicle-to-Everything (V2X) security integration, artificial intelligence and machine learning model security for autonomous systems, zero trust architecture adoption for automotive networks, and development of industry-wide OTA security certification frameworks.")

def create_conclusion(doc):
    """Create conclusion"""
    add_heading_custom(doc, "10. Conclusion and Future Considerations", 1)
    
    paras = [
        "This report has presented a comprehensive automobility cybersecurity strategy for securing Over-the-Air update mechanisms within connected vehicles. The framework integrates technical security controls, ethical considerations, governance processes, and strategic business analysis into a cohesive, academically rigorous, and practically implementable whole. The principal contributions of this research include: (1) a systematic threat model employing the STRIDE methodology and ISO/SAE 21434 TARA framework, identifying ten critical threat categories spanning the OTA ecosystem and mapping them to adversary profiles ranging from opportunistic attackers to nation-state Advanced Persistent Threat (APT) groups; (2) a defense-in-depth security architecture implementing four protective layers ensuring that compromise of any individual control does not result in systemic security failure; (3) an ethical framework grounded in the NIST Privacy Framework and GDPR principles, prioritizing human safety, data privacy, equitable treatment, and transparency; (4) a lifecycle governance model achieving compliance with UNECE WP.29 R155 Cybersecurity Management System requirements and establishing continuous improvement mechanisms; and (5) a comprehensive business case demonstrating positive return on investment through regulatory compliance, avoidance of catastrophic recall costs, competitive market positioning, and enablement of emerging software-defined vehicle revenue models.",
        
        "The analysis demonstrates that securing OTA update systems transcends purely technical considerations, requiring deliberate integration of engineering rigor, ethical responsibility, organizational governance, and economic pragmatism. The automotive industry's transition toward software-defined vehicles will achieve widespread societal acceptance only if the underlying update mechanisms are demonstrably trustworthy, resilient against sophisticated adversaries, and aligned with the safety expectations that society legitimately demands of transportation infrastructure (UNECE, 2020; ISO/SAE, 2021).",
        
        "Future research and development must address several emerging challenges: post-quantum cryptography migration as NIST finalizes standardized quantum-resistant algorithms; Vehicle-to-Everything (V2X) security integration as cooperative intelligent transport systems become prevalent; artificial intelligence and machine learning model security for autonomous driving systems; adoption of zero trust architectural principles within automotive networks; and international harmonization of automotive cybersecurity standards to reduce compliance complexity for global manufacturers. Collaborative defense through expanded participation in the Automotive Information Sharing and Analysis Center (Auto-ISAC) and analogous information-sharing communities will be essential for collective resilience against evolving threat landscapes. The stakes—human safety, individual privacy, economic stability, and the future of sustainable mobility—justify sustained, substantial commitment to automotive cybersecurity excellence."
    ]
    
    for para in paras:
        add_paragraph_custom(doc, para)

def create_references(doc):
    """Create professional references"""
    add_heading_custom(doc, "References", 1)
    
    references = [
        "Auto-ISAC. (2023). Automotive Information Sharing and Analysis Center: Best practices and threat intelligence sharing. Retrieved from https://automotiveisac.com/",
        
        "Barker, E. (2020). Recommendation for key management: Part 1 – General (NIST Special Publication 800-57 Part 1, Revision 5). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-57pt1r5",
        
        "Checkoway, S., McCoy, D., Kantor, B., Anderson, D., Shacham, H., Savage, S., Koscher, K., Czeskis, A., Roesner, F., & Kohno, T. (2011). Comprehensive experimental analyses of automotive attack surfaces. Proceedings of the 20th USENIX Security Symposium, 77–92.",
        
        "Cichonski, P., Millar, T., Grance, T., & Scarfone, K. (2012). Computer security incident handling guide (NIST Special Publication 800-61, Revision 2). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-61r2",
        
        "ENISA. (2019). Cybersecurity challenges in the uptake of artificial intelligence in autonomous driving. European Union Agency for Cybersecurity. Retrieved from https://www.enisa.europa.eu/publications/smart-cars",
        
        "Halder, S., Ozdenizci Kose, B., Edwards, B., & Guvenc, I. (2020). Secure OTA software updates in connected vehicles: A survey. IEEE Access, 8, 220946–220965. https://doi.org/10.1109/ACCESS.2020.3043413",
        
        "ISO/SAE. (2021). ISO/SAE 21434:2021 — Road vehicles — Cybersecurity engineering. International Organization for Standardization. https://www.iso.org/standard/70918.html",
        
        "Koscher, K., Czeskis, A., Roesner, F., Patel, S., Kohno, T., Checkoway, S., McCoy, D., Kantor, B., Anderson, D., Shacham, H., & Savage, S. (2010). Experimental security analysis of a modern automobile. Proceedings of the 2010 IEEE Symposium on Security and Privacy, 447–462. https://doi.org/10.1109/SP.2010.34",
        
        "Lim, H. S. M., Taeihagh, A., & Tan, A. (2022). Governing autonomous vehicles: Emerging responses for safety, liability, privacy, cybersecurity, and industry risks. Transport Reviews, 42(3), 370–395.",
        
        "McKinsey & Company. (2023). The software-defined vehicle: Mapping the automotive software and electronics landscape through 2030. McKinsey & Company. Retrieved from https://www.mckinsey.com/industries/automotive-and-assembly/our-insights/mapping-the-automotive-software-and-electronics-landscape",
        
        "Miller, C., & Valasek, C. (2015). Remote exploitation of an unaltered passenger vehicle. Black Hat USA 2015.",
        
        "NIST. (2018). Platform firmware resiliency guidelines (Special Publication 800-193). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-193",
        
        "NIST. (2020). NIST Privacy Framework: A tool for improving privacy through enterprise risk management (Version 1.0). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.CSWP.01162020",
        
        "NIST. (2022). Post-quantum cryptography: Selected algorithms 2022. National Institute of Standards and Technology. Retrieved from https://csrc.nist.gov/projects/post-quantum-cryptography",
        
        "NIST. (2024). The NIST Cybersecurity Framework (CSF) 2.0 (NIST CSWP 29). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.CSWP.29",
        
        "NTIA. (2021). The minimum elements for a software bill of materials (SBOM). National Telecommunications and Information Administration, U.S. Department of Commerce. Retrieved from https://www.ntia.gov/report/2021/minimum-elements-software-bill-materials-sbom",
        
        "Rescorla, E. (2018). The Transport Layer Security (TLS) Protocol Version 1.3 (RFC 8446). Internet Engineering Task Force. https://doi.org/10.17487/RFC8446",
        
        "UNECE. (2020). UN Regulation No. 155 – Cybersecurity and cybersecurity management system. United Nations Economic Commission for Europe. Retrieved from https://unece.org/transport/documents/2021/03/standards/un-regulation-no-155-cyber-security-and-cyber-security",
        
        "Upstream Security. (2023). Global automotive cybersecurity report. Upstream Security. Retrieved from https://upstream.auto/"
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)

def main():
    # Create document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Configure Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Generate all sections
    create_title_page(doc)
    
    # Add proper TOC
    add_heading_custom(doc, "Table of Contents", 1)
    add_toc(doc)
    doc.add_page_break()
    
    create_executive_summary(doc)
    doc.add_page_break()
    create_introduction(doc)
    doc.add_page_break()
    create_system_description(doc)
    doc.add_page_break()
    create_security_objectives(doc)
    doc.add_page_break()
    
    # Add ethical section
    add_heading_custom(doc, "4. Ethical and Societal Impact Assessment", 1)
    add_paragraph_custom(doc, "The cybersecurity of Over-the-Air update mechanisms carries profound ethical implications given the safety-critical nature of automotive systems, the privacy-sensitive data processed during update operations, and the societal-scale impact of potential security failures. This section examines the ethical dimensions of OTA cybersecurity and delineates how the proposed strategy addresses these concerns, grounded in principles from the NIST Privacy Framework (NIST, 2020) and UNECE WP.29 regulatory requirements (UNECE, 2020).")
    
    add_heading_custom(doc, "4.1 Safety and the Duty of Care", 2)
    add_paragraph_custom(doc, "The preeminent ethical obligation in automotive cybersecurity is the protection of human life. Original Equipment Manufacturers bear a non-delegable duty of care to ensure their products do not expose operators, passengers, or the general public to unreasonable risks of physical harm. When OTA updates modify safety-critical systems—including advanced driver-assistance algorithms, electronic stability control parameters, braking system calibration, or steering assistance functions—the consequences of a cryptographically compromised update could be immediately fatal. The 2015 Jeep Cherokee demonstration established that remote exploitation of vehicle connectivity can result in loss of vehicle control and potential collision (Miller & Valasek, 2015). The proposed strategy addresses this imperative through cryptographic signature verification ensuring only authentic, unmodified firmware executes on safety-critical ECUs; dual-bank A/B partitioning guaranteeing that a cryptographically verified known-good firmware image remains available for emergency fallback; hardware-enforced fail-safe modes that maintain essential vehicle functionality (steering, braking, lighting) even during fault conditions; and phased rollout procedures that limit initial deployment to statistically significant canary populations before fleet-wide release.")
    
    add_heading_custom(doc, "4.2 Privacy and Data Protection", 2)
    add_paragraph_custom(doc, "The OTA update process inherently involves the collection, transmission, and processing of sensitive data including Vehicle Identification Numbers, current firmware version manifests, geolocation data, vehicle telemetry, and potentially biometric information from driver monitoring systems. Inadequate protection of this data enables surveillance, tracking, profiling, and identity theft. The strategy implements the NIST Privacy Framework's core principles (NIST, 2020): data minimization collecting only information strictly necessary for update delivery and verification; pseudonymization and aggregation of telemetry data prior to cloud transmission to prevent individual identification; informed consent mechanisms requiring explicit vehicle owner authorization for data collection where required by the General Data Protection Regulation; encryption of all data in transit via TLS 1.3 with perfect forward secrecy; and retention limitation policies ensuring personal data is maintained only for periods necessary for regulatory compliance or operational requirements.")
    
    add_heading_custom(doc, "4.3 Equity, Transparency, and Public Trust", 2)
    add_paragraph_custom(doc, "Ethical OTA security requires equitable distribution of protective updates without discrimination based on vehicle trim level, geographic location, subscription status, or socioeconomic indicators of the vehicle owner. Security-critical patches must be made available to all affected vehicles regardless of commercial considerations. The strategy mandates universal availability of safety and security updates. Transparency requirements include clear, accessible release notes explaining update contents; publicly available cybersecurity disclosure policies; and timely notification of security incidents affecting vehicle owners. These measures are essential for maintaining societal trust in connected vehicle technologies, which constitutes a prerequisite for widespread adoption of autonomous and connected mobility solutions that promise substantial safety and environmental benefits (Lim et al., 2022; Upstream Security, 2023).")
    
    doc.add_page_break()
    
    create_threat_modeling(doc)
    doc.add_page_break()
    create_security_architecture(doc)
    doc.add_page_break()
    create_process_governance(doc)
    doc.add_page_break()
    create_business_considerations(doc)
    doc.add_page_break()
    create_validation(doc)
    doc.add_page_break()
    create_conclusion(doc)
    doc.add_page_break()
    create_references(doc)
    
    # Save document
    output_path = '/home/siva/.openclaw/workspace/CYB400_OTA_Security_Report_PROFESSIONAL.docx'
    doc.save(output_path)
    print(f"Professional PhD-level report generated successfully: {output_path}")
    print("\nIMPORTANT: To generate the Table of Contents:")
    print("1. Open the document in Microsoft Word")
    print("2. Click on the '[Table of Contents]' placeholder")
    print("3. Right-click and select 'Update Field'")
    print("4. Choose 'Update entire table'")

if __name__ == '__main__':
    main()
